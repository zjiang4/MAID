# -*- coding: utf-8 -*-
from __future__ import annotations

# ==============================================================================
# Core Dependencies Import
# ==============================================================================
import asyncio
import base64
import binascii
import datetime
import functools
import html
import io
import joblib
import os
import random
import re
import time
import traceback
from pathlib import Path
from typing import Any, AsyncGenerator, Dict, List, Optional, Union
from dataclasses import dataclass, field

import numpy as np
import pandas as pd
from scipy.sparse import hstack, csr_matrix
import xgboost as xgb
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.model_selection import train_test_split, GridSearchCV
from sklearn.ensemble import RandomForestRegressor
from sklearn.metrics import mean_absolute_error, mean_squared_error, r2_score
from sentence_transformers import SentenceTransformer

import httpx
from nicegui import app, ui, context, run
from fastapi import FastAPI, HTTPException
from pydantic import BaseModel

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from model_providers import (
    AIModel,
    NVIDIA_DEMO_MODELS,
    assign_healthy_text_models,
    check_nvidia_demo_models,
)
from usmle_step1_syllabus import USMLE_STEP1_SYLLABUS

# --- Image RAG Module ---
try:
    import faiss
    FAISS_AVAILABLE = True
except ImportError:
    FAISS_AVAILABLE = False
    print("Warning: FAISS not installed. Image RAG features will be disabled.")


# ==============================================================================
# Image RAG Module: Classes and State
# ==============================================================================
@dataclass
class ImageRAGState:
    enabled: bool = False
    index: Any = None
    image_descriptions: Dict[int, str] = field(default_factory=dict)
    image_paths: Dict[int, Path] = field(default_factory=dict)
    image_embeddings: np.ndarray | None = None
    embedding_dimension: int = 768
    next_id: int = 0


class ImageRAGManager:
    def __init__(self, base_dir: Path, embedding_dim: int = 768):
        self.base_dir = base_dir
        self.index_file = base_dir / "image_rag_index.faiss"
        self.metadata_file = base_dir / "image_rag_metadata.pkl"
        self.embedding_dim = embedding_dim
        self.state = ImageRAGState(enabled=FAISS_AVAILABLE, embedding_dimension=embedding_dim)
        self._load_index()

    def _load_index(self):
        if not FAISS_AVAILABLE:
            return
        try:
            if self.index_file.exists() and self.metadata_file.exists():
                self.state.index = faiss.read_index(str(self.index_file))
                import pickle
                with open(self.metadata_file, 'rb') as f:
                    metadata = pickle.load(f)
                    self.state.image_descriptions = metadata.get('descriptions', {})
                    self.state.image_paths = {k: Path(v) for k, v in metadata.get('paths', {}).items()}
                    self.state.next_id = metadata.get('next_id', 0)
        except Exception as e:
            print(f"Error loading Image RAG index: {e}")
            self._initialize_index()

    def _initialize_index(self):
        if not FAISS_AVAILABLE:
            return
        self.state.index = faiss.IndexFlatL2(self.embedding_dim)
        self.state.image_descriptions = {}
        self.state.image_paths = {}
        self.state.next_id = 0

    def _save_index(self):
        if not FAISS_AVAILABLE or self.state.index is None:
            return
        try:
            faiss.write_index(self.state.index, str(self.index_file))
            import pickle
            metadata = {
                'descriptions': self.state.image_descriptions,
                'paths': {k: str(v) for k, v in self.state.image_paths.items()},
                'next_id': self.state.next_id
            }
            with open(self.metadata_file, 'wb') as f:
                pickle.dump(metadata, f)
        except Exception as e:
            print(f"Error saving Image RAG index: {e}")

    async def add_image(self, image_path: Path, description: str, embedding_model: Any):
        if not FAISS_AVAILABLE or self.state.index is None:
            return None
        try:
            embedding = embedding_model.encode([description])[0]
            faiss_id = self.state.next_id
            self.state.index.add(np.array([embedding]).astype('float32'))
            self.state.image_descriptions[faiss_id] = description
            self.state.image_paths[faiss_id] = image_path
            self.state.next_id += 1
            self._save_index()
            return faiss_id
        except Exception as e:
            print(f"Error adding image to RAG: {e}")
            return None

    def search_similar_images(self, query_text: str, embedding_model: Any, top_k: int = 5) -> List[Dict]:
        if not FAISS_AVAILABLE or self.state.index is None or self.state.index.ntotal == 0:
            return []
        try:
            query_embedding = embedding_model.encode([query_text])[0]
            query_embedding = np.array([query_embedding]).astype('float32')
            distances, indices = self.state.index.search(query_embedding, min(top_k, self.state.index.ntotal))
            results = []
            for dist, idx in zip(distances[0], indices[0]):
                if idx >= 0 and idx in self.state.image_descriptions:
                    results.append({
                        'id': int(idx),
                        'description': self.state.image_descriptions[idx],
                        'path': self.state.image_paths[idx],
                        'distance': float(dist)
                    })
            return results
        except Exception as e:
            print(f"Error searching images: {e}")
            return []

    def get_image_info(self, faiss_id: int) -> Optional[Dict]:
        if faiss_id in self.state.image_descriptions:
            return {'id': faiss_id, 'description': self.state.image_descriptions[faiss_id], 'path': self.state.image_paths[faiss_id]}
        return None

    def clear_all(self):
        if not FAISS_AVAILABLE:
            return
        self._initialize_index()
        self._save_index()

    def get_stats(self) -> Dict:
        return {'total_images': self.state.index.ntotal if self.state.index else 0, 'enabled': self.state.enabled, 'dimension': self.embedding_dim}


async def index_uploaded_images(image_files: List[str], descriptions: List[str], embedding_model_name: str):
    if not FAISS_AVAILABLE:
        return {'success': False, 'error': 'FAISS is not installed', 'stats': {}}
    if len(image_files) != len(descriptions):
        return {'success': False, 'error': 'Number of images and descriptions must match', 'stats': {}}
    if not image_files:
        return {'success': False, 'error': 'No images provided', 'stats': {}}
    try:
        from sentence_transformers import SentenceTransformer
        embedding_model = SentenceTransformer(embedding_model_name)
        success_count = 0
        errors = []
        for img_path_str, description in zip(image_files, descriptions):
            img_path = Path(img_path_str)
            if not img_path.exists():
                errors.append(f"Image not found: {img_path}")
                continue
            faiss_id = await run.io_bound(IMAGE_RAG_MANAGER.add_image, img_path, description, embedding_model)
            if faiss_id is not None:
                success_count += 1
            else:
                errors.append(f"Failed to index: {img_path.name}")
        stats = IMAGE_RAG_MANAGER.get_stats()
        return {'success': True, 'success_count': success_count, 'total_count': len(image_files), 'errors': errors, 'stats': stats}
    except Exception as e:
        return {'success': False, 'error': str(e), 'stats': {}}


def search_rag_images(query: str, top_k: int = 3) -> List[Dict]:
    if not APP_STATE.get('rag_search_enabled', False):
        return []
    rag_model_name = APP_STATE.get('rag_model_name')
    if not rag_model_name:
        return []
    try:
        from sentence_transformers import SentenceTransformer
        embedding_model = SentenceTransformer(rag_model_name)
        return IMAGE_RAG_MANAGER.search_similar_images(query, embedding_model, top_k)
    except Exception as e:
        print(f"Error in RAG search: {e}")
        return []


def format_rag_results_for_prompt(results: List[Dict]) -> str:
    if not results:
        return "No relevant images found in database."
    formatted = []
    for i, result in enumerate(results, 1):
        img_path = Path(result['path'])
        try:
            relative_path = img_path.relative_to(BASE_OUTPUT_DIR)
            web_url = f"/outputs/{relative_path}"
        except ValueError:
            web_url = str(img_path)
        formatted.append(f"Image {i}: {result['description']}\n  Image URL: {web_url}")
    return "\n".join(formatted)


# ==============================================================================
# Difficulty Prediction Module: API Models, Classes and State
# ==============================================================================
class PredictionItem(BaseModel):
    stem: str
    options: Optional[Dict[str, str]] = None
    answer: Optional[str] = None

class PredictionRequest(BaseModel):
    items: List[Union[str, PredictionItem]]

class PredictionResponse(BaseModel):
    predictions: List[float]
    model_info: str

@dataclass
class ColumnMap:
    id_col: str | None = None
    stem_col: str | None = None
    option_a_col: str | None = None
    option_b_col: str | None = None
    option_c_col: str | None = None
    option_d_col: str | None = None
    option_e_col: str | None = None
    answer_col: str | None = None
    difficulty_col: str | None = None

@dataclass
class AppState:
    df_raw: pd.DataFrame | None = None
    df_preview: pd.DataFrame | None = None
    col_map: ColumnMap = field(default_factory=ColumnMap)
    logs: list[str] = field(default_factory=list)
    df_processed: pd.DataFrame | None = None
    feature_matrix: Any | None = None
    target_vector: pd.Series | None = None
    trained_model: Any | None = None
    validation_results: dict | None = None
    predictions_df: pd.DataFrame | None = None
    feature_method: str | None = None
    preprocess_settings: dict = field(default_factory=dict)
    tfidf_vectorizer: TfidfVectorizer | None = None
    stat_feature_names: list[str] = field(default_factory=list)
    sentence_model_name: str | None = None
    sentence_model: Any | None = None


# ==============================================================================
# MCQ Development Module: Global Configuration and State
# ==============================================================================
BASE_OUTPUT_DIR = Path("mcq_outputs")
IMAGES_DIR = BASE_OUTPUT_DIR / "generated_images"
SELECTION_LOGS_DIR = BASE_OUTPUT_DIR / "selection_logs"
IMAGE_RAG_DIR = BASE_OUTPUT_DIR / "image_rag"

SYLLABUS_DATA_HIERARCHICAL = USMLE_STEP1_SYLLABUS

PREDEFINED_KEYS = {}

APP_STATE: Dict[str, Any] = {
    "saved_models_config": [],
    "selected_kps_with_inputs": {},
    "keyword_to_details_map": {},
    "is_generating": False,
    "stop_requested": False,
    "image_rag": None,
    "rag_search_enabled": False,
    "rag_model_name": None,
    "difficulty_predictor": AppState(),
}

IMAGE_RAG_MANAGER: ImageRAGManager = None  # type: ignore

class StopRequestedError(Exception):
    pass

# ==============================================================================
# Difficulty Prediction Module: Core Background Functions
# ==============================================================================
def log_difficulty_predictor(msg: str):
    ts = time.strftime('%H:%M:%S')
    log_msg = f"[{ts}] {msg}"
    APP_STATE["difficulty_predictor"].logs.insert(0, log_msg)

def _perform_encoding(texts: list[str], model_name_val: str) -> np.ndarray:
    predictor_state = APP_STATE["difficulty_predictor"]
    if predictor_state.sentence_model is None or predictor_state.sentence_model_name != model_name_val:
        log_difficulty_predictor(f"Loading sentence model: {model_name_val}...")
        predictor_state.sentence_model = SentenceTransformer(model_name_val)
        predictor_state.sentence_model_name = model_name_val
    return predictor_state.sentence_model.encode(texts)

def _perform_training(X_train: Any, y_train: Any, model_config: dict) -> Any:
    model_type = model_config['type']
    if model_type == 'RandomForest':
        model = RandomForestRegressor(n_estimators=200, random_state=42, n_jobs=-1)
        model.fit(X_train, y_train)
        return model
    elif model_type == 'XGBoost':
        model = xgb.XGBRegressor(n_estimators=200, learning_rate=0.1, max_depth=6, random_state=42, n_jobs=-1, verbosity=0)
        if model_config.get('autotune'):
            param_grid = {'n_estimators': [100, 200], 'max_depth': [4, 6, 8], 'learning_rate': [0.05, 0.1]}
            model = GridSearchCV(model, param_grid, cv=3, scoring='neg_mean_absolute_error', n_jobs=-1)
            model.fit(X_train, y_train)
        else:
            model.fit(X_train, y_train)
        return model
    return None

# ==============================================================================
# MCQ Development Module: Core Classes
# ==============================================================================
class MCQDevelopmentSystem:
    def __init__(self, models_config_list: List[Dict]):
        self.models = []
        self.history: List[Dict] = []
        for c in models_config_list:
            try:
                self.models.append(AIModel.from_config(c))
            except Exception as e:
                print(f"Warning: Skipping invalid model config '{c.get('name', 'N/A')}': {e}")

    def select_model(self, role: str, specified_model_name: Optional[str] = None) -> AIModel:
        if role in ['Image Generation', 'Multimodal']:
            image_models = [m for m in self.models if m.model_type in ['Image Generation', 'Multimodal']]
            if specified_model_name:
                model = next((m for m in image_models if m.name == specified_model_name), None)
                if model:
                    return model
            if image_models:
                return image_models[0]
        else:
            text_models = [m for m in self.models if m.model_type == 'Text Output']
            if specified_model_name:
                model = next((m for m in text_models if m.name == specified_model_name), None)
                if model:
                    return model
            if text_models:
                return text_models[0]
        raise ValueError(f"No suitable model found for role '{role}'.")

    def add_to_history(self, role: str, content: str, version: int = 1, usage: Optional[Dict] = None):
        entry = {
            "timestamp": datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "role": role, "content": content, "version": version
        }
        self.history.append(entry)

    def get_plain_text_history(self) -> str:
        return "\n".join([f"[{e['role']}]: {e['content']}" for e in self.history])

    def clear_history(self):
        self.history = []


# ==============================================================================
# Common Helper Functions
# ==============================================================================
def _normalize_colons(text: str) -> str:
    return text.replace('\uff1a', ':').replace('\uff1b', ':') if isinstance(text, str) else ""


def save_content_to_file(content: str, folder: Path, base_name: str, extension: str) -> Path:
    folder.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    file_path = folder / f"{base_name}_{timestamp}.{extension}"
    file_path.write_text(content, encoding='utf-8')
    return file_path


def convert_to_tree_nodes(hierarchical_data: dict, parent_key: str = '') -> List[Dict]:
    nodes = []
    for key, value in hierarchical_data.items():
        node_id = f"{parent_key}/{key}" if parent_key else key
        node = {'id': node_id, 'label': key}
        if isinstance(value, dict):
            node['children'] = convert_to_tree_nodes(value, node_id)
        nodes.append(node)
    return nodes


def get_nodes_as_dict(nodes: List[Dict]) -> Dict[str, Dict]:
    node_dict = {}
    for node in nodes:
        node_dict[node['id']] = node
        if 'children' in node:
            node_dict.update(get_nodes_as_dict(node['children']))
    return node_dict


def flatten_syllabus_to_map(hierarchical_data: dict, parent_path: Optional[List[str]] = None) -> Dict[str, Dict]:
    result = {}
    for key, value in hierarchical_data.items():
        current_path = (parent_path or []) + [key]
        if isinstance(value, dict):
            result.update(flatten_syllabus_to_map(value, current_path))
        elif isinstance(value, list):
            for item in value:
                item_path = current_path + [item]
                result[item] = {
                    "topic": current_path[-3] if len(current_path) > 2 else "General",
                    "unit": current_path[-2] if len(current_path) > 1 else "General",
                    "item": item,
                    "point": item,
                    "requirement": "None"
                }
    return result


def parse_markdown_syllabus(md_text: str) -> Dict:
    lines = md_text.split('\n')
    root: Dict[str, Any] = {}
    heading_stack: List[tuple] = []
    pending_items: List[str] = []

    def flush_items(target_dict: Dict):
        nonlocal pending_items
        if pending_items and target_dict is not None:
            for item in pending_items:
                if item not in target_dict:
                    target_dict[item] = []
            pending_items = []

    for line in lines:
        line = line.rstrip()
        hm = re.match(r'^(#{1,6})\s+(.+)$', line)
        if hm:
            level = len(hm.group(1))
            title = hm.group(2).strip()
            current_target = heading_stack[-1][1] if heading_stack else root
            flush_items(current_target)
            while heading_stack and heading_stack[-1][0] >= level:
                heading_stack.pop()
            new_node: Dict[str, Any] = {}
            if heading_stack:
                _, parent_dict = heading_stack[-1]
                parent_dict[title] = new_node
            else:
                root[title] = new_node
            heading_stack.append((level, new_node))
        elif line.strip().startswith(('- ', '* ')):
            item = re.sub(r'^[\s\-\*]+', '', line).strip()
            if item:
                pending_items.append(item)
    if heading_stack:
        flush_items(heading_stack[-1][1])
    else:
        flush_items(root)
    return root


def normalize_syllabus_dict(data: Any) -> Dict:
    if not isinstance(data, dict):
        return {}
    result = {}
    for key, value in data.items():
        if isinstance(value, dict):
            result[key] = normalize_syllabus_dict(value)
        elif isinstance(value, list):
            result[key] = [v if isinstance(v, str) else str(v) for v in value]
        elif isinstance(value, str):
            result[key] = [value]
        else:
            result[key] = value
    return result


def generate_writer_prompt_context(keyword: str, requirement: str, knowledge_map: Dict) -> Dict:
    details = knowledge_map.get(keyword)
    if not details:
        return {"topic": "General", "unit": "General", "item": "Basic Concept", "point": keyword, "requirement": requirement}
    return details


def generate_html_report_content(history: List[Dict], title: str, prompt_context: Dict, image_path: Optional[Path] = None) -> str:
    html_content = f"""<!DOCTYPE html><html><head><meta charset='utf-8'><title>{html.escape(title)}</title>
<style>body{{font-family:Arial,sans-serif;max-width:900px;margin:20px auto;padding:20px;}}
.qa{{border:1px solid #ddd;border-radius:8px;padding:20px;margin-bottom:15px;}}
.role{{font-weight:bold;color:#2c3e50;margin-bottom:10px;padding:5px 10px;background:#f0f0f0;border-radius:4px;display:inline-block;}}
.version{{font-size:0.85em;color:#888;}}
.generated-image{{max-width:400px;border-radius:8px;margin:10px 0;}}</style></head><body>
<h1>MCQ Development Report: {html.escape(title)}</h1>"""
    for entry in history:
        role = entry.get('role', 'Unknown')
        content = entry.get('content', '')
        version = entry.get('version', 1)
        html_content += f'<div class="qa"><span class="role">{html.escape(role)}</span> <span class="version">v{version}</span><pre style="white-space:pre-wrap;font-family:inherit;">{html.escape(content)}</pre></div>'
    if image_path and image_path.exists():
        try:
            image_data = base64.b64encode(image_path.read_bytes()).decode('utf-8')
            html_content += f"<div class='qa'><h3>Generated Image</h3><img src='data:image/png;base64,{image_data}' class='generated-image'></div>"
        except Exception as e:
            html_content += f"<p style='color:red;'>Image embed failed: {html.escape(str(e))}</p>"
    html_content += '</div></body></html>'
    return html_content


def extract_final_question_from_history(history: List[Dict]) -> Optional[str]:
    if not history:
        return None
    def _find_question_block(content: str) -> Optional[str]:
        normalized = _normalize_colons(content)
        if not normalized:
            return None
        match = re.search(r"Question:.*", normalized, re.DOTALL | re.IGNORECASE)
        if match:
            qt = match.group(0).strip()
            if all(k in qt for k in ["Options:", "Correct Answer:", "Explanation:"]):
                return qt
        return None
    final_decision = next((e.get('content', '') for e in reversed(history) if "Final Decision" in e.get('role', '')), None)
    if final_decision:
        normalized = _normalize_colons(final_decision)
        text_match = re.search(r"\[Final Question Text\]:\s*(.*)", normalized, re.DOTALL | re.IGNORECASE)
        if text_match and (q := _find_question_block(text_match.group(1))):
            return q
        if (q := _find_question_block(normalized)):
            return q
    for entry in reversed(history):
        if (q := _find_question_block(entry.get('content', ''))):
            return q
    return None


async def generate_and_save_image(image_model: AIModel, prompt: str, output_dir: Path) -> tuple:
    try:
        ui.notify("Image generation: URL mode started...", spinner=True)
        response_text, _ = await image_model.call_ai_model(
            "You are an image generation assistant. Return ONLY a valid URL to the generated image, nothing else.",
            prompt
        )
        url_match = re.search(r'https?://[^\s\'"<>]+', response_text)
        if url_match:
            image_url = url_match.group(0)
            async with httpx.AsyncClient(timeout=60) as client:
                img_resp = await client.get(image_url)
                if img_resp.status_code == 200:
                    output_dir.mkdir(parents=True, exist_ok=True)
                    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
                    save_path = output_dir / f"mcq_image_{timestamp}.png"
                    save_path.write_bytes(img_resp.content)
                    return f"Image downloaded from URL successfully: {image_url}", save_path
            ui.notify(f"Download failed ({image_url}), retrying with Base64 in 2s...", color='warning')
            await asyncio.sleep(2)
    except Exception as e:
        print(f"URL mode failed: {e}")
    try:
        ui.notify("Image generation: Base64 mode started...", spinner=True)
        response_text, _ = await image_model.call_ai_model(
            "You are an image generation assistant. Generate the image described by the user. Return the image as a base64-encoded PNG data URL (data:image/png;base64,...). Return ONLY the data URL.",
            prompt
        )
        b64_match = re.search(r'data:image/[^;]+;base64,([A-Za-z0-9+/=]+)', response_text)
        if b64_match:
            b64_data = b64_match.group(1)
            padding = 4 - (len(b64_data) % 4)
            if padding < 4:
                b64_data += '=' * padding
            image_bytes = base64.b64decode(b64_data)
            output_dir.mkdir(parents=True, exist_ok=True)
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            save_path = output_dir / f"mcq_image_{timestamp}.png"
            save_path.write_bytes(image_bytes)
            return "Image generated via Base64 successfully.", save_path
        return "No valid image data returned by the API.", None
    except Exception as e:
        return f"Base64 mode also failed: {e}", None


def save_paper_as_docx(questions: List[Dict], run_folder: Path, base_name: str, include_answers: bool = True) -> Path:
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Pt(36)
        section.bottom_margin = Pt(36)
    title = doc.add_heading('Medical Item Set', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if include_answers:
        note = doc.add_paragraph('(With Answers and Explanations)')
        note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        note = doc.add_paragraph('(Questions Only)')
        note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    for q_data in questions:
        q_text = _normalize_colons(q_data.get('text', ''))
        question_match = re.search(r"Question:(.*?)(?=Options:)", q_text, re.DOTALL | re.IGNORECASE)
        options_match = re.search(r"Options:(.*?)(?=Correct Answer:)", q_text, re.DOTALL | re.IGNORECASE)
        answer_match = re.search(r"Correct Answer:(.*?)(?=Explanation:)", q_text, re.DOTALL | re.IGNORECASE)
        explanation_match = re.search(r"Explanation:(.*)", q_text, re.DOTALL | re.IGNORECASE)
        stem = question_match.group(1).strip() if question_match else q_text
        p = doc.add_paragraph(f"{q_data['number']}. {stem}")
        p.paragraph_format.space_after = Pt(6)
        if options_match:
            for opt_line in options_match.group(1).strip().split('\n'):
                opt_line = opt_line.strip()
                if opt_line:
                    p = doc.add_paragraph(opt_line)
                    p.paragraph_format.left_indent = Pt(36)
                    p.paragraph_format.space_after = Pt(2)
        if include_answers:
            if answer_match:
                p = doc.add_paragraph(f"Answer: {answer_match.group(1).strip()}")
                p.paragraph_format.space_before = Pt(6)
                p.runs[0].bold = True if p.runs else None
            if explanation_match:
                p = doc.add_paragraph(f"Explanation: {explanation_match.group(1).strip()}")
                p.paragraph_format.space_after = Pt(12)
        if q_data.get('image_path') and Path(q_data['image_path']).exists():
            try:
                doc.add_picture(str(q_data['image_path']), width=Inches(4.0))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                pass
        doc.add_page_break()
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    suffix = "complete" if include_answers else "questions"
    file_path = run_folder / f"{base_name}_{suffix}_{timestamp}.docx"
    doc.save(str(file_path))
    return file_path


WRITER_SYSTEM_PROMPT = """You are a senior clinical medicine expert and medical educator, specializing in designing high-quality Multiple Choice Questions (MCQ) for the Clinical Practitioner Qualification Examination.

The exam assesses clinical doctors' professional knowledge and clinical reasoning abilities after completing their training.

Requirements for generated questions:
1. Questions should be clear, concise, and unambiguous.
2. Options should be highly deceptive, requiring careful analysis.
3. Only one option is the best answer.
4. Avoid options like "none of the above" or "all of the above".

For A2 type (case analysis):
- Build a complete clinical case with patient info, history, physical exam, and investigations.
- Template: [Sex], [Age]. [Chief complaint] + [Duration]. [History]. [Exam findings]. [Investigations]. [Key question]

For A1 type (knowledge):
- Focus on core concepts, mechanisms, classification standards.
- Even without a case, the question must have depth.

Please strictly follow this format for your output:

[Question Type]: A1 or A2, with brief reasoning.

Question:
[Your question stem here]

Options:
[Option text 1]
[Option text 2]
[Option text 3]
[Option text 4]
[Option text 5]

Correct Answer:
[Full text of correct option]

Explanation:
[Detailed explanation of why the correct answer is right and why distractors are wrong]

[Image Generation Prompt]:
[If A2 type and image would help, provide a detailed image generation prompt. Otherwise write "None"]
"""

WRITER_USER_PROMPT_TEMPLATE = """Please create ONE high-quality single-choice MCQ for the Clinical Practitioner Qualification Examination, based on the following content specification:

Content Specification: {point}
Topic: {topic}
Unit: {unit}

{image_instruction}"""

WRITER_USER_PROMPT_WITH_IMAGE = """Please create ONE high-quality single-choice MCQ for the Clinical Practitioner Qualification Examination, based on the following content specification:

Content Specification: {point}
Topic: {topic}
Unit: {unit}

IMPORTANT: An Image Generation Model is configured for this session.
- If this is an A2 question and a visual material (ECG, X-ray, pathology slide, etc.) would significantly enhance the question, please design a question that includes an image.
- Do NOT describe the image features in the question stem.
- Provide a detailed, professional image generation prompt below.

Please strictly follow this format for your output:

[Question Type]: A1 or A2, with brief reasoning.

Question:
[Your question stem here]

Options:
[Option text 1]
[Option text 2]
[Option text 3]
[Option text 4]
[Option text 5]

Correct Answer:
[Full text of correct option]

Explanation:
[Detailed explanation]

[Image Generation Prompt]:
[A detailed prompt for AI image generation, or "None"]"""

REVIEWER_SYSTEM_PROMPT = """You are a senior MCQ reviewer for the Clinical Practitioner Qualification Examination. You are an expert in core clinical specialties and deeply understand how to evaluate clinical reasoning through high-quality MCQs.

Your task is to provide strict, critical feedback on the provided MCQ draft to ensure it meets examination standards. Be constructive and provide specific improvement suggestions.

Review Focus: {review_focus}

Please provide your detailed review comments, then confirm with: "Review complete."
"""

REVIEWER_USER_PROMPT_TEMPLATE = """Please review the following MCQ draft for the Clinical Practitioner Qualification Examination:

{question_text}

Please provide your review comments and modification suggestions. If no issues found, state so clearly."""

EDITOR_SYSTEM_PROMPT = """You are the chief editor for the Clinical Practitioner Qualification Examination MCQ review process.

Your core responsibility is to integrate feedback from multiple reviewers and guide the question writer to make precise revisions.

The final goal is to ensure every MCQ meets the examination's academic standards, difficulty, and discrimination.

Your output MUST include these sections:
[Final Conclusion]: (One-sentence final assessment)
[Core Issues Summary]: (1-2 sentences summarizing key issues)
[Revision Guide]: (Specific, actionable revision guidance)

If reviewer opinions conflict, weigh them and give clear recommendations."""

FINAL_DECISION_SYSTEM_PROMPT = """You are the final review editor for the Clinical Practitioner Qualification Examination.

Your task is to review the entire MCQ development history (draft, reviews, editor guidance, and writer's revision) and make a final decision.

Your output MUST include:
[Final Decision]: Choose one: "Approved for Exam" / "Approved with Minor Revisions" / "Requires Major Revisions"
[Review Comments]: Detailed reasoning for your decision.
[Final Question Text]: The complete, ready-to-use question text (if approved)."""


async def run_full_mcq_pipeline(mcq_system: MCQDevelopmentSystem, prompt_context: Dict, model_names: Dict, step_progress, step_label, q_num: int, is_image_gen_enabled: bool) -> tuple:
    history = []
    usage_stats = {
        "Writer_Initial": {"usage": None, "time": 0.0, "model": ""},
        "Image_Generation": {"usage": None, "time": 0.0, "model": ""},
        "Reviewer_1": {"usage": None, "time": 0.0, "model": ""},
        "Reviewer_2": {"usage": None, "time": 0.0, "model": ""},
        "Reviewer_3": {"usage": None, "time": 0.0, "model": ""},
        "Editor_Summary": {"usage": None, "time": 0.0, "model": ""},
        "Writer_Revision": {"usage": None, "time": 0.0, "model": ""},
        "Editor_FinalDecision": {"usage": None, "time": 0.0, "model": ""},
    }
    log_prefix = f"Q{q_num}_"
    current_step = 0
    total_steps = 8 if is_image_gen_enabled else 7

    def update_progress(step: int, label: str):
        nonlocal current_step
        current_step = step
        step_progress.value = step / total_steps
        step_label.text = f"Q{q_num}: {label} ({step}/{total_steps})"

    # Step 1: Writer generates initial question
    update_progress(1, "Writer drafting...")
    start_time = time.time()
    writer_ai = mcq_system.select_model("Writer", model_names["writer"])
    if is_image_gen_enabled:
        user_prompt = WRITER_USER_PROMPT_WITH_IMAGE.format(point=prompt_context["point"], topic=prompt_context["topic"], unit=prompt_context["unit"], image_instruction="Image generation is enabled.")
    else:
        user_prompt = WRITER_USER_PROMPT_TEMPLATE.format(point=prompt_context["point"], topic=prompt_context["topic"], unit=prompt_context["unit"], image_instruction="Image generation is NOT enabled. Focus on text-only questions.")
    question_text, usage = await mcq_system.call_ai_model(writer_ai, WRITER_SYSTEM_PROMPT, user_prompt)
    usage_stats["Writer_Initial"].update({"usage": usage, "time": time.time() - start_time, "model": writer_ai.name})
    mcq_system.add_to_history(f"{log_prefix}Writer ({writer_ai.name})", question_text, 1, usage)

    # Step 2: Image generation (if enabled)
    q_img_path = None
    q_img_prompt = None
    if is_image_gen_enabled:
        update_progress(2, "Image generation...")
        start_time = time.time()
        try:
            image_model = mcq_system.select_model("Image Generation", model_names.get("image_generator"))
            prompt_match = re.search(r"\[Image Generation Prompt\]:\s*(.*)", question_text, re.DOTALL | re.IGNORECASE)
            image_prompt_text = prompt_match.group(1).strip() if prompt_match else None
            if image_prompt_text and image_prompt_text.lower() != "none" and len(image_prompt_text) > 10:
                image_log, q_img_path = await generate_and_save_image(image_model, image_prompt_text, IMAGES_DIR)
                q_img_prompt = image_prompt_text
                mcq_system.add_to_history(f"{log_prefix}Image Generation ({image_model.name})", image_log, 2)
            else:
                mcq_system.add_to_history(f"{log_prefix}Image Generation", "Writer did not provide a valid image prompt.", 2)
            usage_stats["Image_Generation"].update({"time": time.time() - start_time, "model": image_model.name if 'image_model' in dir() else ""})
        except Exception as e:
            mcq_system.add_to_history(f"{log_prefix}Image Generation - Failed", str(e), 2)
    else:
        q_img_path = None

    # Steps 3-5: Three reviewers
    review_focuses = [
        "Medical content accuracy and clinical relevance",
        "Question structure and technical compliance",
        "Language clarity and formatting"
    ]
    v_offset = 1 if is_image_gen_enabled else 0
    for i, focus in enumerate(review_focuses):
        update_progress(3 + i, f"Reviewer {i+1}...")
        start_time = time.time()
        reviewer_ai = mcq_system.select_model(f"Reviewer_{i+1}", model_names["reviewers"][i])
        rev_user_prompt = f"{REVIEWER_USER_PROMPT_TEMPLATE}\n\nReview Focus: {focus}\n\n{question_text}"
        review, usage = await mcq_system.call_ai_model(reviewer_ai, REVIEWER_SYSTEM_PROMPT.format(review_focus=focus), rev_user_prompt)
        usage_stats[f"Reviewer_{i+1}"].update({"usage": usage, "time": time.time() - start_time, "model": reviewer_ai.name})
        mcq_system.add_to_history(f"{log_prefix}Reviewer {i+1} ({reviewer_ai.name})", review, 3 + v_offset + i, usage)

    # Step 6: Editor summarizes
    update_progress(6, "Editor summary...")
    start_time = time.time()
    editor_ai = mcq_system.select_model("Editor", model_names["editor"])
    reviews_text = mcq_system.get_plain_text_history()
    editor_user_prompt = f"""Please integrate the following reviews for a Clinical Practitioner Qualification Examination MCQ:

{reviews_text}

Based on all reviews, provide a clear, concise, actionable revision guide for the original writer.
Focus on improving: clinical knowledge depth, clinical reasoning assessment, and distractor quality."""
    summary, usage = await mcq_system.call_ai_model(editor_ai, EDITOR_SYSTEM_PROMPT, editor_user_prompt)
    usage_stats["Editor_Summary"].update({"usage": usage, "time": time.time() - start_time, "model": editor_ai.name})
    mcq_system.add_to_history(f"{log_prefix}Editor Summary ({editor_ai.name})", summary, 6 + v_offset, usage)

    # Step 7: Writer revises
    update_progress(7, "Writer revision...")
    start_time = time.time()
    revision_prompt = f"""You are the original writer of this Clinical Practitioner Qualification Examination MCQ.
Please carefully read the editor's summary and revision guide, and revise your initial draft.
Provide the complete revised MCQ with all required sections (Question, Options, Correct Answer, Explanation).

{summary}

{question_text}"""
    revision, usage = await mcq_system.call_ai_model(writer_ai, WRITER_SYSTEM_PROMPT, revision_prompt)
    usage_stats["Writer_Revision"].update({"usage": usage, "time": time.time() - start_time, "model": writer_ai.name})
    mcq_system.add_to_history(f"{log_prefix}Writer Revision ({writer_ai.name})", revision, 7 + v_offset, usage)

    # Step 8: Final decision
    update_progress(8, "Final decision...")
    start_time = time.time()
    final_prompt = f"""As the final review editor, please review this MCQ's entire development process.

{mcq_system.get_plain_text_history()}

Make a final decision on the revised MCQ. Your output must contain:
[Final Decision]: Approved for Exam / Approved with Minor Revisions / Requires Major Revisions
[Review Comments]: Detailed reasoning
[Final Question Text]: The complete question text"""
    decision, usage = await mcq_system.call_ai_model(editor_ai, FINAL_DECISION_SYSTEM_PROMPT, final_prompt)
    usage_stats["Editor_FinalDecision"].update({"usage": usage, "time": time.time() - start_time, "model": editor_ai.name})
    mcq_system.add_to_history(f"{log_prefix}Final Decision ({editor_ai.name})", decision, 8 + v_offset, usage)

    return history, usage_stats, q_img_path, q_img_prompt


# ==============================================================================
# Difficulty Prediction Module: UI Components
# ==============================================================================
def create_difficulty_predictor_ui(panels_container, upload_tab, preprocess_tab, features_tab, model_tab, predict_tab, api_tab, ui_refs: Dict):
    STATE = APP_STATE["difficulty_predictor"]

    with panels_container:
        with ui.tab_panel(upload_tab).classes('p-4'):
            ui.label('Data Upload & Column Mapping').classes('text-2xl font-bold')
            with ui.card().classes('w-full'):
                ui.label('1) Upload File').classes('text-lg font-semibold')
                ui.label('Supports CSV (UTF-8) or Excel (.xlsx).').classes('text-gray-500')
                upload_card = ui.card_section()
                def handle_upload(e):
                    try:
                        bytes_io = io.BytesIO(e.content.read())
                        df = pd.read_csv(bytes_io) if e.name.lower().endswith('.csv') else pd.read_excel(bytes_io)
                        STATE.df_raw = df
                        log_difficulty_predictor(f"File uploaded: '{e.name}'")
                        ui.notify('File uploaded successfully!', type='positive')
                        preview_container.clear()
                        with preview_container:
                            ui.table(columns=[{'name': c, 'label': c, 'field': c, 'sortable': True} for c in df.columns], rows=df.head(5).to_dict('records')).props('flat bordered pagination=5 rows-per-page')
                    except Exception as ex:
                        log_difficulty_predictor(f"File parse error: {ex}")
                        ui.notify(f"File parse error: {ex}", type='negative')
                ui.upload(on_upload=handle_upload, auto_upload=True).props('accept=.csv,.xlsx,.xls').classes('mt-2')
            with ui.card().classes('w-full mt-4'):
                ui.label('2) Column Mapping').classes('text-lg font-semibold')
                ui.label('Select the stem (question text) and difficulty columns.').classes('text-gray-500')
                preview_container = ui.column().classes('w-full')
                with preview_container:
                    ui.label('Please upload data first').classes('text-gray-500')
                def upload_card():
                    pass
                def confirm_mapping():
                    STATE.df_processed = STATE.df_raw.copy()
                    log_difficulty_predictor("Column mapping confirmed, generating preview.")
                    preview_container.clear()
                    with preview_container:
                        if STATE.df_processed is not None:
                            ui.table(columns=[{'name': c, 'label': c, 'field': c, 'sortable': True} for c in STATE.df_processed.columns], rows=STATE.df_processed.head(5).to_dict('records')).props('flat bordered')
                            ui.notify("Mapping confirmed.", type='positive')
                ui.button('Confirm Mapping', on_click=confirm_mapping, icon='check').classes('mt-4')

        with ui.tab_panel(preprocess_tab).classes('p-4'):
            ui.label('Text Preprocessing').classes('text-2xl font-bold')
            with ui.card().classes('w-full'):
                lower_switch = ui.switch('Convert to lowercase')
                punct_switch = ui.switch('Remove punctuation')
                async def handle_preprocess():
                    if STATE.df_processed is None:
                        return ui.notify('Please upload data and complete column mapping first.', type='negative')
                    STATE.preprocess_settings = {'lower': lower_switch.value, 'punct': punct_switch.value}
                    ui.notify('Preprocessing applied.', type='positive')
                ui.button('Apply Preprocessing', on_click=handle_preprocess).classes('mt-4')

        with ui.tab_panel(features_tab).classes('p-4'):
            ui.label('Feature Engineering').classes('text-2xl font-bold')
            ui.label('Select stem and difficulty columns first.').classes('text-gray-500')
            with ui.row().classes('w-full gap-4 mt-4'):
                with ui.card().classes('w-1/2'):
                    stem_col_select = ui.select(list(STATE.df_processed.columns) if STATE.df_processed is not None else [], label='Stem Column')
                    diff_col_select = ui.select(list(STATE.df_processed.columns) if STATE.df_processed is not None else [], label='Difficulty Column')
                    async def handle_tfidf():
                        if STATE.df_processed is None:
                            return ui.notify('Please upload and map data first.', type='warning')
                        stem_col = stem_col_select.value
                        diff_col = diff_col_select.value
                        if not stem_col or not diff_col:
                            return ui.notify('Please select both stem and difficulty columns.', type='warning')
                        STATE.col_map.stem_col = stem_col
                        STATE.col_map.difficulty_col = diff_col
                        texts = STATE.df_processed[stem_col].fillna('').astype(str).tolist()
                        STATE.tfidf_vectorizer = TfidfVectorizer(max_features=5000)
                        STATE.feature_matrix = STATE.tfidf_vectorizer.fit_transform(texts)
                        STATE.target_vector = STATE.df_processed[diff_col].astype(float)
                        STATE.feature_method = 'tfidf'
                        log_difficulty_predictor(f"TF-IDF features generated, shape: {STATE.feature_matrix.shape}")
                        ui.notify('TF-IDF features generated successfully!', type='positive')
                    ui.button('Generate TF-IDF Features', on_click=handle_tfidf)
                with ui.card().classes('w-1/2'):
                    model_select = ui.select(['shibing624/text2vec-base-chinese', 'all-MiniLM-L6-v2', 'all-mpnet-base-v2'], label='Sentence Transformer Model')
                    async def handle_embeddings():
                        if STATE.df_processed is None:
                            return ui.notify('Please upload and map data first.', type='warning')
                        stem_col = stem_col_select.value
                        diff_col = diff_col_select.value
                        if not stem_col or not diff_col:
                            return ui.notify('Please select both stem and difficulty columns.', type='warning')
                        STATE.col_map.stem_col = stem_col
                        STATE.col_map.difficulty_col = diff_col
                        texts = STATE.df_processed[stem_col].fillna('').astype(str).tolist()
                        STATE.feature_matrix = await run.cpu_bound(_perform_encoding, texts, model_select.value)
                        STATE.target_vector = STATE.df_processed[diff_col].astype(float)
                        STATE.feature_method = 'embedding'
                        log_difficulty_predictor(f"Sentence embeddings generated, shape: {STATE.feature_matrix.shape}")
                        ui.notify('Sentence embeddings generated successfully!', type='positive')
                    ui.button('Generate Sentence Embeddings', on_click=handle_embeddings)
            with ui.card().classes('w-full mt-4'):
                ui.label('Statistical Features').classes('text-lg font-semibold')
                stat_cols_select = ui.select(list(STATE.df_processed.columns) if STATE.df_processed is not None else [], label='Text columns', multiple=True)
                features_to_add = {
                    'char_count': ui.checkbox('Character Count'),
                    'word_count': ui.checkbox('Word Count'),
                    'sentence_count': ui.checkbox('Sentence Count'),
                    'avg_word_length': ui.checkbox('Avg Word Length'),
                }
                def handle_stat_features():
                    if STATE.feature_matrix is None:
                        return ui.notify('Please generate baseline features (TF-IDF or embeddings) first.', type='warning')
                    if not stat_cols_select.value:
                        return ui.notify('Please select at least one text column.', type='warning')
                    selected_features = [key for key, checkbox in features_to_add.items() if checkbox.value]
                    if not selected_features:
                        return ui.notify('Please select at least one feature type.', type='warning')
                    df = STATE.df_processed.copy()
                    df['stat_combined_text'] = df[stat_cols_select.value].fillna('').astype(str).agg(' '.join, axis=1)
                    new_arrays, new_names = [], []
                    if 'char_count' in selected_features:
                        new_arrays.append(df['stat_combined_text'].str.len().values.reshape(-1, 1))
                        new_names.append('char_count')
                    if 'word_count' in selected_features:
                        new_arrays.append(df['stat_combined_text'].str.split().str.len().values.reshape(-1, 1))
                        new_names.append('word_count')
                    if 'sentence_count' in selected_features:
                        count = df['stat_combined_text'].str.split(r'[.!?]').apply(lambda x: len([s for s in x if s.strip()]))
                        new_arrays.append(count.values.reshape(-1, 1))
                        new_names.append('sentence_count')
                    if 'avg_word_length' in selected_features:
                        char_c = df['stat_combined_text'].str.len()
                        word_c = df['stat_combined_text'].str.split().str.len()
                        avg_l = (char_c / word_c.replace(0, 1)).fillna(0)
                        new_arrays.append(avg_l.values.reshape(-1, 1))
                        new_names.append('avg_word_length')
                    if not new_arrays:
                        return
                    stat_features = np.hstack(new_arrays)
                    if STATE.stat_feature_names is None:
                        STATE.stat_feature_names = []
                    unique_new = [n for n in new_names if n not in STATE.stat_feature_names]
                    STATE.stat_feature_names.extend(unique_new)
                    STATE.feature_matrix = hstack([STATE.feature_matrix, stat_features]).tocsr() if isinstance(STATE.feature_matrix, csr_matrix) else np.hstack([STATE.feature_matrix, stat_features])
                    log_difficulty_predictor(f"Statistical features appended: {', '.join(new_names)}, new shape: {STATE.feature_matrix.shape}")
                    ui.notify(f"Statistical features appended! Added {len(new_names)} features.", type='positive')
                ui.button('Append Statistical Features', on_click=handle_stat_features).classes('mt-4')

    with panels_container:
        with ui.tab_panel(model_tab).classes('p-4'):
            ui.label('Model Training & Evaluation').classes('text-2xl font-bold')
            @ui.refreshable
            def build_results_display():
                if STATE.validation_results is None:
                    with ui.card().classes('w-full'):
                        ui.label('Training results will appear here').classes('text-gray-500')
                    return
                res = STATE.validation_results
                with ui.card().classes('w-full'):
                    ui.label('Validation Results').classes('text-lg font-semibold')
                    ui.markdown(f"- **Model**: {res['model_name']}\n- **MAE**: {res['mae']:.4f}\n- **RMSE**: {res['rmse']:.4f}\n- **R2**: {res['r2']:.4f}")
                with ui.card().classes('w-full'):
                    ui.label('Full Prediction Results').classes('text-lg font-semibold')
                    df = STATE.predictions_df.copy()
                    df['error'] = abs(df[STATE.col_map.difficulty_col] - df['predictions'])
                    cols_to_show = [c for c in [STATE.col_map.stem_col, STATE.col_map.difficulty_col, 'predictions', 'error'] if c and c in df.columns]
                    ui.table(columns=[{'name': c, 'label': c, 'field': c, 'sortable': True, 'align': 'left'} for c in cols_to_show], rows=df.to_dict('records'), pagination=10).props('flat bordered')
                with ui.card().classes('w-full'):
                    def export_predictions():
                        ui.download(STATE.predictions_df.to_csv(index=False).encode('utf-8-sig'), 'predictions.csv')
                    def export_model():
                        buffer = io.BytesIO()
                        joblib.dump({'model': STATE.trained_model, 'feature_method': STATE.feature_method, 'tfidf_vectorizer': STATE.tfidf_vectorizer, 'sentence_model_name': STATE.sentence_model_name, 'stat_feature_names': STATE.stat_feature_names, 'preprocess_settings': STATE.preprocess_settings}, buffer)
                        ui.download(buffer.getvalue(), 'difficulty_pipeline.joblib')
                    with ui.row():
                        ui.button('Export Predictions (CSV)', on_click=export_predictions)
                        ui.button('Export Model (.joblib)', on_click=export_model)
            with ui.card().classes('w-full mt-4'):
                ui.label('Model Training').classes('text-lg font-semibold')
                model_type = ui.select(['RandomForest', 'XGBoost'], value='RandomForest', label='Select Model')
                with ui.row().classes('w-full items-center'):
                    ui.label('Test Set Ratio').classes('text-base')
                    test_size = ui.slider(min=0.1, max=0.5, step=0.05, value=0.2).classes('flex-grow mx-4')
                    ui.number(label='Exact Value', min=0.1, max=0.5, step=0.05, format='%.2f').classes('w-28').bind_value(test_size, 'value')
                autotune_switch = ui.switch('Auto-tune (XGBoost, very slow)').bind_visibility_from(model_type, 'value', value='XGBoost')
                spinner_container = ui.row()
                async def handle_train():
                    if STATE.feature_matrix is None:
                        return ui.notify('Please generate features first.', type='warning')
                    with spinner_container:
                        ui.spinner(size='lg', color='primary')
                    X_train, X_test, y_train, y_test = train_test_split(STATE.feature_matrix, STATE.target_vector, test_size=test_size.value, random_state=42)
                    model_config = {'type': model_type.value, 'autotune': autotune_switch.value}
                    model = await run.cpu_bound(_perform_training, X_train, y_train, model_config)
                    spinner_container.clear()
                    STATE.trained_model = model.best_estimator_ if isinstance(model, GridSearchCV) else model
                    if isinstance(model, GridSearchCV):
                        log_difficulty_predictor(f"GridSearch best params: {model.best_params_}")
                    y_pred = STATE.trained_model.predict(X_test)
                    STATE.validation_results = {'model_name': model_type.value, 'mae': mean_absolute_error(y_test, y_pred), 'rmse': np.sqrt(mean_squared_error(y_test, y_pred)), 'r2': r2_score(y_test, y_pred)}
                    full_preds = STATE.trained_model.predict(STATE.feature_matrix)
                    df = STATE.df_processed.copy()
                    df['predictions'] = full_preds
                    STATE.predictions_df = df
                    log_difficulty_predictor(f"Training complete. R2: {STATE.validation_results['r2']:.4f}")
                    ui.notify('Model training complete!', type='positive')
                    build_results_display.refresh()
                ui.button('Start Training', on_click=handle_train, icon='play_circle_outline')
            build_results_display()

    with panels_container:
        with ui.tab_panel(predict_tab).classes('p-4'):
            ui.label('Difficulty Prediction').classes('text-2xl font-bold')
            with ui.row().classes('w-full gap-4 mt-4'):
                with ui.card().classes('w-1/2'):
                    ui.label('Single Question Prediction').classes('text-lg font-semibold')
                    question_text = ui.textarea(label='Paste question text here', placeholder='e.g., A 65-year-old male presents with...').props('outlined autogrow').classes('w-full')
                    spinner_single = ui.row()
                    result_single = ui.column().classes('w-full')
                    async def handle_single_prediction():
                        if not STATE.trained_model:
                            return ui.notify('No trained model yet.', type='negative')
                        if not question_text.value:
                            return ui.notify('Please enter question text.', type='warning')
                        with spinner_single:
                            ui.spinner(size='lg', color='primary')
                        text = question_text.value
                        if STATE.preprocess_settings.get('lower'):
                            text = text.lower()
                        if STATE.preprocess_settings.get('punct'):
                            text = re.sub(r'[^\w\s]', ' ', text)
                        if STATE.feature_method == 'tfidf':
                            feature_vec = STATE.tfidf_vectorizer.transform([text])
                        else:
                            feature_vec = await run.cpu_bound(_perform_encoding, [text], STATE.sentence_model_name)
                        if STATE.stat_feature_names:
                            stat_vals = []
                            ts = pd.Series([text])
                            if 'char_count' in STATE.stat_feature_names:
                                stat_vals.append(ts.str.len())
                            if 'word_count' in STATE.stat_feature_names:
                                stat_vals.append(ts.str.split().str.len())
                            if 'sentence_count' in STATE.stat_feature_names:
                                stat_vals.append(ts.str.split(r'[.!?]').apply(lambda x: len([s for s in x if s.strip()])))
                            if 'avg_word_length' in STATE.stat_feature_names:
                                stat_vals.append((ts.str.len() / ts.str.split().str.len().replace(0, 1)).fillna(0))
                            stat_features = pd.concat(stat_vals, axis=1).values
                            feature_vec = hstack([feature_vec, stat_features]).tocsr() if isinstance(feature_vec, csr_matrix) else np.hstack([feature_vec, stat_features])
                        prediction = STATE.trained_model.predict(feature_vec)
                        spinner_single.clear()
                        result_single.clear()
                        with result_single:
                            with ui.card().classes('w-full items-center'):
                                ui.label('Predicted Difficulty').classes('font-semibold')
                                ui.label(f'{prediction[0]:.2f}').classes('text-5xl font-bold text-blue-600 mt-2')
                    ui.button('Predict Difficulty', on_click=handle_single_prediction, icon='lightbulb').classes('w-full mt-4')
                with ui.card().classes('w-1/2'):
                    ui.label('Batch Prediction').classes('text-lg font-semibold')
                    ui.label('Upload CSV or Excel with question texts.').classes('text-gray-500')
                    batch_state = {'df': None, 'name': None}
                    def handle_batch_upload(e):
                        try:
                            bytes_io = io.BytesIO(e.content.read())
                            df = pd.read_csv(bytes_io) if e.name.lower().endswith('.csv') else pd.read_excel(bytes_io)
                            batch_state['df'] = df
                            batch_state['name'] = e.name
                            ui.notify(f"File '{e.name}' uploaded.", type='positive')
                            mapping_container.clear()
                            with mapping_container:
                                text_col_sel = ui.select(list(df.columns), label='Select text column').classes('w-full')
                                ui.button('Start Batch Prediction', on_click=lambda: handle_batch_predict(text_col_sel.value), icon='play_circle_outline')
                        except Exception as ex:
                            ui.notify(f"File parse error: {ex}", type='negative')
                    ui.upload(on_upload=handle_batch_upload, auto_upload=True).props('accept=.csv,.xlsx,.xls').classes('mt-2')
                    mapping_container = ui.column().classes('w-full mt-4 gap-4')
                    result_batch = ui.column().classes('w-full mt-4')
                    async def handle_batch_predict(text_col):
                        if batch_state['df'] is None or text_col is None or not STATE.trained_model:
                            return ui.notify('Please ensure file uploaded, column selected, and model trained.', type='negative')
                        result_batch.clear()
                        with result_batch:
                            ui.spinner(size='md', color='primary')
                        try:
                            df = batch_state['df']
                            texts = df[text_col].astype(str).tolist()
                            processed = [re.sub(r'[^\w\s]', ' ', t.lower()) if STATE.preprocess_settings.get('punct') and STATE.preprocess_settings.get('lower') else t for t in texts]
                            if STATE.feature_method == 'tfidf':
                                fm = STATE.tfidf_vectorizer.transform(processed)
                            else:
                                fm = await run.cpu_bound(_perform_encoding, processed, STATE.sentence_model_name)
                            if STATE.stat_feature_names:
                                sdf = pd.DataFrame({'text': processed})
                                sv = []
                                if 'char_count' in STATE.stat_feature_names:
                                    sv.append(sdf['text'].str.len())
                                if 'word_count' in STATE.stat_feature_names:
                                    sv.append(sdf['text'].str.split().str.len())
                                if 'sentence_count' in STATE.stat_feature_names:
                                    sv.append(sdf['text'].str.split(r'[.!?]').apply(lambda x: len([s for s in x if s.strip()])))
                                if 'avg_word_length' in STATE.stat_feature_names:
                                    sv.append((sdf['text'].str.len() / sdf['text'].str.split().str.len().replace(0, 1)).fillna(0))
                                sf = pd.concat(sv, axis=1).values
                                fm = hstack([fm, sf]).tocsr() if isinstance(fm, csr_matrix) else np.hstack([fm, sf])
                            preds = STATE.trained_model.predict(fm)
                            result_df = df.copy()
                            result_df['predicted_difficulty'] = preds
                            result_batch.clear()
                            with result_batch:
                                ui.label('Prediction complete!').classes('font-semibold')
                                ui.table(columns=[{'name': c, 'label': c, 'field': c, 'align': 'left'} for c in result_df.columns], rows=result_df.head(10).to_dict('records')).props('flat bordered')
                                def download_results():
                                    ui.download(result_df.to_csv(index=False, encoding='utf-8-sig').encode(), f'predictions_{batch_state["name"]}')
                                ui.button('Download Full Results', on_click=download_results, icon='download')
                        except Exception as e:
                            ui.notify(f"Prediction error: {e}", type='negative')
                            result_batch.clear()

    with panels_container:
        with ui.tab_panel(api_tab).classes('p-4'):
            ui.label('System Integration & Logs').classes('text-2xl font-bold text-gray-800')
            with ui.row().classes('w-full gap-4'):
                with ui.column().classes('w-2/3 gap-4'):
                    with ui.card().classes('w-full'):
                        ui.label('API Documentation').classes('text-lg font-semibold')
                        ui.markdown("The system provides REST API endpoints for integrating question difficulty prediction into external applications. **Please train a model in the UI before using the API.**")
                    with ui.card().classes('w-full'):
                        ui.label('1. Check Model Status').classes('text-lg font-semibold')
                        ui.badge('GET', color='positive')
                        ui.markdown("Check if a trained prediction model is available.")
                        with ui.element('div').classes('p-2 bg-slate-100 rounded-md font-mono text-sm w-full overflow-x-auto mt-2'):
                            ui.code("/api/status")
                        ui.label('curl example:').classes('font-bold mt-4')
                        with ui.element('div').classes('p-2 bg-slate-100 rounded-md font-mono text-sm w-full overflow-x-auto'):
                            ui.code('curl "http://127.0.0.1:8080/api/status"')
                        ui.label('Example response:').classes('font-bold mt-4')
                        with ui.element('div').classes('p-2 bg-slate-100 rounded-md font-mono text-sm w-full overflow-x-auto'):
                            ui.code('{"model_trained": true, "feature_method": "embedding", "timestamp": 1678886400.0}')
                    with ui.card().classes('w-full'):
                        ui.label('2. Predict Question Difficulty').classes('text-lg font-semibold')
                        ui.badge('POST', color='info')
                        ui.markdown("Send one or more questions. Returns predicted difficulty values.")
                        with ui.element('div').classes('p-2 bg-slate-100 rounded-md font-mono text-sm w-full overflow-x-auto mt-2'):
                            ui.code("/api/predict")
                        with ui.expansion('Simple Usage (text only)', icon='code').classes('w-full mt-4'):
                            ui.markdown("Each item can be a simple string.")
                            with ui.element('div').classes('p-2 bg-slate-100 rounded-md font-mono text-sm w-full overflow-x-auto'):
                                ui.code('''curl -X POST "http://127.0.0.1:8080/api/predict" -H "Content-Type: application/json" -d '{"items": ["Question text here..."]}' ''')
                        with ui.expansion('Advanced Usage (structured)', icon='data_object').classes('w-full'):
                            ui.markdown("Each item can be a JSON object with stem, options, answer.")
                            with ui.element('div').classes('p-2 bg-slate-100 rounded-md font-mono text-sm w-full overflow-x-auto'):
                                ui.code('''{"items": [{"stem": "Question text...", "options": {"A": "Opt A", "B": "Opt B"}, "answer": "Opt A"}]}''')
                with ui.card().classes('w-1/3'):
                    with ui.row().classes('w-full justify-between items-center'):
                        ui.label('Activity Logs').classes('text-xl font-semibold text-gray-700')
                        ui.button(icon='delete_sweep', on_click=lambda: STATE.logs.clear()).classes('text-sm').tooltip('Clear all logs')
                    log_view = ui.log(max_lines=100).classes('w-full h-80 mt-2 bg-white border')
                    def update_logs():
                        log_view.clear()
                        if not STATE.logs:
                            log_view.push('[INFO] No logs yet.')
                        else:
                            for item in reversed(STATE.logs):
                                if "[ERROR]" in item:
                                    log_view.push(f'<span style="color: red;">{item}</span>')
                                elif "[WARN]" in item:
                                    log_view.push(f'<span style="color: orange;">{item}</span>')
                                else:
                                    log_view.push(item)
                    ui.timer(2.0, update_logs)
                    update_logs()


# ==============================================================================
# UI Layout: Main Page
# ==============================================================================
@app.get('/outputs')
def serve_outputs():
    pass

@ui.page('/')
def main_page():
    app.add_static_files('/outputs', BASE_OUTPUT_DIR)
    ui.add_head_html("""<style>.soft-card{border-radius:12px;box-shadow:0 4px 16px rgba(0,0,0,.05);}</style>""")
    tree_nodes = convert_to_tree_nodes(SYLLABUS_DATA_HIERARCHICAL)
    ALL_NODES_BY_ID = get_nodes_as_dict(tree_nodes)
    APP_STATE["keyword_to_details_map"] = flatten_syllabus_to_map(SYLLABUS_DATA_HIERARCHICAL)

    ui_refs: Dict[str, Any] = {
        "model_selects": {}, "model_table": None, "status_text": None,
        "question_preview_area": None, "download_row": None, "progress_row": None,
        "step_progress": None, "step_label": None, "generated_questions": [],
        "run_folder": None, "log_file_path": None, "generate_button": None, "stop_button": None
    }

    def render_question_preview_area(questions: List[Dict]):
        qpa = ui_refs['question_preview_area']
        qpa.clear()
        with qpa:
            if not questions:
                ui.label("No questions generated.").classes('m-4 text-gray-500')
                return
            for q_data in sorted(questions, key=lambda x: x['number']):
                card = ui.card().classes('w-full mb-2 cursor-pointer hover:bg-gray-100 transition-colors')
                card.on('click', lambda q=q_data, c=card: open_edit_dialog(c, q))
                with card:
                    if img_path_str := q_data.get('image_path'):
                        if Path(img_path_str).exists():
                            relative_path = Path(img_path_str).relative_to(BASE_OUTPUT_DIR)
                            ui.image(f"/outputs/{relative_path}?t={time.time()}").classes('w-full max-w-sm rounded-md mx-auto my-2')
                    cleaned = re.sub(r'^\s*Question:\s*', '', _normalize_colons(q_data['text']), flags=re.IGNORECASE).strip()
                    escaped = html.escape(cleaned).replace('\n', '<br>')
                    title_html = f"<b>{q_data['number']}.</b> {escaped}"
                    if 'predicted_difficulty' in q_data:
                        ds = q_data['predicted_difficulty']
                        color = 'green' if ds < 0.4 else '#E69500' if ds < 0.7 else 'red'
                        title_html += f' <br><span style="color:{color};font-style:italic;font-size:0.9em;">(Difficulty: {ds:.2f})</span>'
                    ui.html(title_html).style('white-space: pre-wrap;')

    async def handle_predict_all_difficulties():
        STATE = APP_STATE["difficulty_predictor"]
        if not STATE.trained_model:
            ui.notify('Please train a model in Difficulty Model Training tab first.', type='negative', position='center', timeout=5000)
            return
        all_q = ui_refs.get('generated_questions')
        if not all_q:
            ui.notify('No generated questions to predict.', type='warning')
            return
        notification = ui.notification('Predicting difficulty for all questions...', spinner=True, timeout=None)
        try:
            texts = []
            for q_data in all_q:
                full = _normalize_colons(q_data['text'])
                texts.append(full)
            processed = [re.sub(r'[^\w\s]', ' ', t.lower()) if STATE.preprocess_settings.get('punct') and STATE.preprocess_settings.get('lower') else t for t in texts]
            if STATE.feature_method == 'tfidf':
                fm = STATE.tfidf_vectorizer.transform(processed)
            else:
                fm = await run.cpu_bound(_perform_encoding, processed, STATE.sentence_model_name)
            if STATE.stat_feature_names:
                sdf = pd.DataFrame({'text': processed})
                sv = []
                if 'char_count' in STATE.stat_feature_names:
                    sv.append(sdf['text'].str.len())
                if 'word_count' in STATE.stat_feature_names:
                    sv.append(sdf['text'].str.split().str.len())
                if 'sentence_count' in STATE.stat_feature_names:
                    sv.append(sdf['text'].str.split(r'[.!?]').apply(lambda x: len([s for s in x if s.strip()])))
                if 'avg_word_length' in STATE.stat_feature_names:
                    sv.append((sdf['text'].str.len() / sdf['text'].str.split().str.len().replace(0, 1)).fillna(0))
                sf = pd.concat(sv, axis=1).values
                fm = hstack([fm, sf]).tocsr() if isinstance(fm, csr_matrix) else np.hstack([fm, sf])
            preds = STATE.trained_model.predict(fm)
            for i, q_data in enumerate(all_q):
                q_data['predicted_difficulty'] = preds[i]
            render_question_preview_area(all_q)
            update_questions()
            notification.message = f"Successfully predicted difficulty for {len(all_q)} questions!"
            notification.spinner = False
            await asyncio.sleep(4)
        except Exception as e:
            log_difficulty_predictor(f"Batch prediction error: {e}")
            ui.notify(f"Prediction error: {e}", type='negative', multi_line=True)
        finally:
            notification.dismiss()

    def create_all_download_buttons(paper_full: Path, paper_questions: Path, log_file: Path):
        download_row = ui_refs.get('download_row')
        if not download_row:
            return
        download_row.clear()
        with download_row:
            with ui.grid(columns=2).classes('w-full gap-2'):
                ui.button('Download Item Set (DOCX)', on_click=lambda: ui.download(str(paper_questions)), icon='article')
                ui.button('Download Full Set with Answers (DOCX)', on_click=lambda: ui.download(str(paper_full)), icon='description')
                ui.button('Download Generation Log (TXT)', on_click=lambda: ui.download(str(log_file)), icon='history_edu')
                ui.button('Predict Question Difficulty', on_click=handle_predict_all_difficulties, icon='analytics', color='secondary').tooltip('Predict difficulty for all generated questions')
        download_row.set_visibility(True)

    def update_questions():
        run_folder = ui_refs.get('run_folder')
        all_q = ui_refs.get('generated_questions')
        log_file_path = ui_refs.get('log_file_path')
        if not all([run_folder, all_q, log_file_path]):
            ui.notify("Cannot update: missing runtime info.", color='warning')
            return
        try:
            paper_full = save_paper_as_docx(all_q, run_folder, "item_set", include_answers=True)
            paper_questions = save_paper_as_docx(all_q, run_folder, "item_set", include_answers=False)
            create_all_download_buttons(paper_full, paper_questions, log_file_path)
            ui.notify("Paper files updated!", color='positive')
        except Exception as e:
            print(f"Update DOCX error: {e}")
            ui.notify(f"Update DOCX failed: {e}", color='negative', multi_line=True)

    async def open_edit_dialog(card_to_update: ui.card, question_data: dict):
        chat_history = question_data.setdefault('chat_history', [])
        chat_ui_map: Dict[str, ui.element] = {}
        if not chat_history:
            initial = question_data['text']
            if question_data.get('image_path'):
                initial += f"\n\n[Current Image Prompt]: {question_data.get('image_prompt', 'None')}"
            chat_history.append({'role': 'system_context', 'text': initial, 'name': 'Original Question', 'sent': False, 'id': f'msg_{time.time()}'})

        is_a2 = bool(re.search(r"Question:\s*(男|女),", _normalize_colons(question_data.get('text', '')), re.IGNORECASE))

        with ui.dialog() as dialog, ui.card().style('width: 90vw; max-width: 90vw; height: 90vh; max-height: 90vh;').classes('flex flex-col no-wrap'):
            with ui.row().classes('w-full justify-between items-center flex-shrink-0 p-2 bg-gray-100 rounded-t-lg'):
                ui.markdown(f"### Edit Question {question_data['number']}")
                with ui.row().classes('gap-2'):
                    def export_chat():
                        lines = [f"# Question {question_data['number']} - Chat History\n"]
                        for msg in chat_history:
                            lines.append(f"\n---\n[{msg.get('name', 'System')}]\n{msg.get('text', '')}")
                        filename = f"Q{question_data['number']}_chat_{datetime.datetime.now():%Y%m%d%H%M}.txt"
                        ui.download('\n'.join(lines).encode('utf-8'), filename)
                    ui.button("Export Chat", on_click=export_chat, icon='download', color='secondary')
                    ui.button("Save Changes", on_click=save_changes, icon='save', color='positive')
                    ui.button("Close", on_click=dialog.close, icon='close')

            with ui.row().classes('w-full flex-grow no-wrap overflow-hidden'):
                with ui.card().classes('w-3/5 h-full flex flex-col no-wrap'):
                    all_models = sorted([c.get('name', 'unknown') for c in APP_STATE["saved_models_config"]])
                    text_models = sorted([c['name'] for c in APP_STATE["saved_models_config"] if c.get('model_type') == 'Text Output'])
                    model_select = ui.select(all_models, label='Select Chat Model', value=text_models[0] if text_models else None).classes('w-full p-2 flex-shrink-0')
                    with ui.element('div').classes('w-full flex-grow overflow-y-auto p-2 border-t border-b') as chat_container:
                        for msg in chat_history:
                            if 'id' not in msg:
                                msg['id'] = f'msg_{time.time()}'
                            with ui.chat_message(name=msg['name'], sent=msg['sent']).style('white-space: pre-wrap; background-color: #f0f0f0;' if msg['role'] == 'system_context' else '') as msg_ui:
                                chat_ui_map[msg['id']] = msg_ui
                                with ui.row().classes('w-full items-start justify-between no-wrap'):
                                    if msg.get('is_image_update'):
                                        ui.image(msg['image_url']).classes('w-full max-w-xs rounded-md')
                                        ui.html(f'<pre style="white-space:pre-wrap;font-family:inherit;">{html.escape(msg["text"])}</pre>').classes('flex-grow')
                                    else:
                                        ui.html(f'<pre style="white-space:pre-wrap;font-family:inherit;">{html.escape(msg["text"])}</pre>').classes('flex-grow')
                                    if msg.get('role') == 'user':
                                        ui.button(icon='delete_outline', on_click=lambda m=msg: handle_delete(m)).props('flat round dense size=sm').classes('ml-2')

                    async def send_message(user_text: Optional[str] = None, actual_prompt: Optional[str] = None):
                        prompt_for_api = ""
                        if actual_prompt and not user_text:
                            text_to_display = actual_prompt
                            prompt_for_api = actual_prompt
                        else:
                            text_to_display = user_text if user_text else input_field.value.strip()
                            if not text_to_display:
                                return
                            input_field.value = ''
                            prompt_for_api = actual_prompt if actual_prompt else text_to_display
                        if user_text or not actual_prompt:
                            new_msg = {'role': 'user', 'text': text_to_display, 'name': 'Me', 'sent': True, 'id': f'msg_{time.time()}'}
                            chat_history.append(new_msg)
                            with chat_container:
                                with ui.chat_message(name='Me', sent=True) as new_ui:
                                    chat_ui_map[new_msg['id']] = new_ui
                                    with ui.row().classes('w-full items-start justify-between no-wrap'):
                                        ui.markdown(text_to_display).style('white-space: pre-wrap;').classes('flex-grow')
                                        ui.button(icon='delete_outline', on_click=lambda m=new_msg: handle_delete(m)).props('flat round dense size=sm').classes('ml-2')
                                await ui.run_javascript(f'getElement({chat_container.id}).scrollTop = getElement({chat_container.id}).scrollHeight')
                        selected_name = model_select.value
                        if not selected_name:
                            ui.notify("Please select a model.", color='negative')
                            return
                        model_config = next((c for c in APP_STATE["saved_models_config"] if c['name'] == selected_name), None)
                        if not model_config:
                            ui.notify(f"Model '{selected_name}' not found.", color='negative')
                            return
                        with chat_container:
                            with ui.chat_message(name='AI', sent=False) as msg_container:
                                ui.spinner(size='md', color='primary')
                            await ui.run_javascript(f'getElement({chat_container.id}).scrollTop = getElement({chat_container.id}).scrollHeight')
                        try:
                            m_type = model_config.get('model_type', 'Text Output')
                            if m_type in ['Image Generation', 'Multimodal']:
                                text_configs = [c for c in APP_STATE["saved_models_config"] if c.get('model_type') == 'Text Output']
                                if not text_configs:
                                    raise ValueError("No text model configured for processing image instructions.")
                                msg_container.clear()
                                with msg_container:
                                    ui.html("<i>AI understanding your image modification request...</i>")
                                text_sys = MCQDevelopmentSystem([text_configs[0]])
                                text_model = text_sys.models[0]
                                orig_prompt = question_data.get('image_prompt', 'None')
                                refine_sys = "You are a medical image prompt engineer. Modify the image prompt based on user request. Output ONLY the new prompt text."
                                refine_user = f"[Original Question]\n{question_data.get('text', 'No info.')}\n\n[User Request]\n'{prompt_for_api}'\n\n[Original Prompt]\n'{orig_prompt}'\n\nProvide the new image prompt:"
                                new_prompt, _ = await text_sys.call_ai_model(text_model, refine_sys, refine_user)
                                new_prompt = new_prompt.strip().strip('"').strip("'")
                                msg_container.clear()
                                with msg_container:
                                    ui.html("<i>Calling image model...</i>")
                                img_sys = MCQDevelopmentSystem([model_config])
                                img_model = img_sys.models[0]
                                img_log, img_path = await generate_and_save_image(img_model, new_prompt, IMAGES_DIR)
                                msg_container.clear()
                                with msg_container:
                                    if img_path:
                                        ui.markdown("**New image generated. Use it?**")
                                        ui.image(f"/outputs/{img_path.relative_to(BASE_OUTPUT_DIR)}?t={time.time()}").classes('w-full max-w-sm')
                                        ui.markdown(f"**New prompt:**\n`{new_prompt}`").classes('text-xs mt-2')
                                        with ui.row() as btn_row:
                                            def accept_img():
                                                question_data['image_path'] = str(img_path)
                                                question_data['image_prompt'] = new_prompt
                                                btn_row.clear()
                                                with btn_row:
                                                    ui.markdown("Image updated.").classes("text-green-600 font-bold")
                                                update_right_panel(right_panel)
                                                ui.notify("Image updated.", color='positive')
                                                img_url = f"/outputs/{img_path.relative_to(BASE_OUTPUT_DIR)}?t={time.time()}"
                                                chat_history.append({'id': f'msg_{time.time()}', 'role': 'assistant', 'name': 'AI', 'is_image_update': True, 'image_url': img_url, 'text': "Image updated.", 'sent': False})
                                                with chat_container:
                                                    with ui.chat_message(name='AI', sent=False) as ru:
                                                        chat_ui_map[ru.id] = ru
                                                        ui.image(img_url).classes('w-full max-w-xs rounded-md')
                                                        ui.markdown("Image updated.")
                                            def discard_img():
                                                Path(img_path).unlink(missing_ok=True)
                                                btn_row.clear()
                                                with btn_row:
                                                    ui.markdown("Cancelled.").classes("text-gray-500")
                                                ui.notify("Cancelled, image deleted.", color='info')
                                            ui.button("Accept", on_click=accept_img, color='positive')
                                            ui.button("Discard", on_click=discard_img, color='negative')
                                    else:
                                        ui.markdown(f"**Image generation failed:**\n\n{img_log}")
                                chat_history.append({'role': 'assistant', 'text': f"(Image operation) {img_log}", 'name': 'AI', 'sent': False})
                            else:
                                dummy_sys = MCQDevelopmentSystem([model_config])
                                m_inst = dummy_sys.models[0]
                                conv = "\n".join([f"{'User' if m['sent'] else 'AI'}: {m['text']}" for m in chat_history[:-1]])
                                sys_prompt = "You are a senior medical exam question expert. Modify questions per user instructions. When providing a revised question, include 'Question:', 'Options:', 'Correct Answer:', and 'Explanation:' so the system can save it."
                                user_prompt = f"[Chat History]\n{conv}\n\n[Latest Request]\n{prompt_for_api}\n\nRespond to the latest request."
                                full_response = ""
                                first_chunk = True
                                async def scroll_bottom():
                                    await ui.run_javascript(f'getElement({chat_container.id}).scrollTop = getElement({chat_container.id}).scrollHeight')
                                stream = dummy_sys.call_ai_model_stream(m_inst, sys_prompt, user_prompt)
                                async for piece in stream:
                                    if first_chunk:
                                        msg_container.clear()
                                        with msg_container:
                                            live = ui.html(f'<pre style="white-space:pre-wrap;font-family:inherit;">{html.escape(piece)}</pre>')
                                        first_chunk = False
                                    full_response += piece
                                    safe = html.escape(full_response)
                                    if live:
                                        live.set_content(f'<pre style="white-space:pre-wrap;font-family:inherit;">{safe}</pre>')
                                    await scroll_bottom()
                                if not first_chunk:
                                    ai_msg = {'role': 'assistant', 'text': full_response, 'name': 'AI', 'sent': False, 'id': f'msg_{time.time()}'}
                                    chat_history.append(ai_msg)
                                    chat_ui_map[ai_msg['id']] = msg_container
                                else:
                                    msg_container.clear()
                                    with msg_container:
                                        ui.markdown("No response from model.")
                        except Exception as e:
                            msg_container.clear()
                            with msg_container:
                                ui.notify(f"Error: {e}", color='negative', multi_line=True)

                    async def handle_delete(msg_to_del: dict):
                        try:
                            idx = chat_history.index(msg_to_del)
                            if (idx + 1) < len(chat_history) and chat_history[idx + 1].get('role') == 'assistant':
                                with ui.dialog() as cd, ui.card():
                                    ui.label('Delete this message and the AI reply too?')
                                    with ui.row().classes('w-full justify-end gap-2'):
                                        ui.button('Delete Both', on_click=lambda: cd.submit('both'), color='negative')
                                        ui.button('Only This', on_click=lambda: cd.submit('single'))
                                        ui.button('Cancel', on_click=lambda: cd.submit('cancel'), color='secondary')
                                result = await cd
                                if result == 'both':
                                    ai_msg = chat_history[idx + 1]
                                    chat_ui_map.pop(ai_msg.get('id'), None).delete()
                                    chat_history.remove(ai_msg)
                                    chat_ui_map.pop(msg_to_del.get('id'), None).delete()
                                    chat_history.remove(msg_to_del)
                                    ui.notify("Both messages deleted.", color='positive')
                                elif result == 'single':
                                    chat_ui_map.pop(msg_to_del.get('id'), None).delete()
                                    chat_history.remove(msg_to_del)
                                    ui.notify("Message deleted.", color='positive')
                                else:
                                    ui.notify("Cancelled.", color='info')
                        except ValueError:
                            ui.notify("Message not found.", color='negative')

                    def update_right_preview_panel(panel: ui.column):
                        try:
                            panel.clear()
                            with panel:
                                ui.markdown("**Original Question**").classes('text-lg font-bold mb-2')
                                if img_str := question_data.get('image_path'):
                                    img_p = Path(img_str)
                                    if img_p.exists():
                                        cache_url = f"/outputs/{img_p.relative_to(BASE_OUTPUT_DIR)}?t={time.time()}"
                                        ui.image(cache_url).classes('w-full max-w-md rounded-md mx-auto my-2')
                                        ui.separator()
                                preview_text = re.sub(r"(\r\n|\n)\[?Image.*\]?:.*", "", question_data.get('text', ''), flags=re.DOTALL | re.IGNORECASE).strip()
                                ui.markdown(preview_text).style('white-space: pre-wrap;')
                        except Exception as e:
                            ui.notify("Preview update failed.", color='negative')

                    def save_changes():
                        last_ai = next((m['text'] for m in reversed(chat_history) if m.get('role') == 'assistant'), None)
                        was_updated = False
                        if last_ai:
                            normalized = _normalize_colons(last_ai)
                            q_match = re.search(r"Question:.*", normalized, re.DOTALL | re.IGNORECASE)
                            if q_match and all(k in normalized for k in ["Options:", "Correct Answer:", "Explanation:"]):
                                question_data['text'] = q_match.group(0).strip()
                                was_updated = True
                                ui.notify("Updated from AI response.", color='positive')
                        card_to_update.clear()
                        with card_to_update:
                            if img_str := question_data.get('image_path'):
                                if Path(img_str).exists():
                                    rel = Path(img_str).relative_to(BASE_OUTPUT_DIR)
                                    ui.image(f"/outputs/{rel}?t={time.time()}").classes('w-full max-w-sm rounded-md mx-auto my-2')
                            cleaned = re.sub(r'^\s*Question:\s*', '', _normalize_colons(question_data['text']), flags=re.IGNORECASE).strip()
                            if 'predicted_difficulty' in question_data:
                                ds = question_data['predicted_difficulty']
                                color = 'green' if ds < 0.4 else '#E69500' if ds < 0.7 else 'red'
                                ui.html(f"<b>{question_data['number']}.</b> {html.escape(cleaned)} <br><span style='color:{color};font-style:italic;font-size:0.9em;'>(Difficulty: {ds:.2f})</span>").style('white-space: pre-wrap;')
                            else:
                                ui.markdown(f"**{question_data['number']}.** {cleaned}").style('white-space: pre-wrap;')
                        try:
                            run_folder = ui_refs.get('run_folder')
                            if not run_folder or not Path(run_folder).exists():
                                run_folder = BASE_OUTPUT_DIR / f"Edited_Session_{datetime.datetime.now():%Y%m%d_%H%M%S}"
                                run_folder.mkdir(parents=True, exist_ok=True)
                                ui_refs['run_folder'] = run_folder
                            log_file_path = ui_refs.get('log_file_path')
                            if not log_file_path or not Path(log_file_path).exists():
                                log_file_path = run_folder / "edit_session_log.txt"
                                log_file_path.write_text(f"Edit session started at {datetime.datetime.now()}\n")
                                ui_refs['log_file_path'] = log_file_path
                            with open(log_file_path, 'a', encoding='utf-8') as f:
                                f.write(f"\n--- Question {question_data['number']} saved at {datetime.datetime.now()} ---\n")
                                f.write("Source: AI text response.\n" if was_updated else "Source: Existing data.\n")
                                f.write(f"New data: {question_data}\n")
                            update_questions()
                        except Exception as e:
                            ui.notify(f"Save error: {e}", color='negative', multi_line=True)
                        final = f"Question {question_data['number']} saved."
                        if not was_updated:
                            final += " (No text update applied)"
                        ui.notify(final, color='positive')
                        dialog.close()

                    with ui.row().classes('w-full items-center p-2 flex-shrink-0'):
                        input_field = ui.textarea(placeholder='Enter modification instructions (Ctrl+Enter to send)...').props('outlined autogrow').classes('flex-grow').on('keydown.enter.ctrl', send_message)
                        ui.button(icon='send', on_click=send_message).props('flat round')

                    with ui.row().classes('w-full justify-start gap-2 p-2 flex-shrink-0') as text_presets:
                        async def handle_preset_click(face_text: str, prompt_tmpl: str):
                            actual = prompt_tmpl.format(question=question_data.get('text', 'No text'))
                            await send_message(user_facing_text=face_text, actual_prompt=actual)
                        ui.button("Optimize Options", on_click=lambda: handle_preset_click("Optimize options.", "Review the following MCQ and optimize its distractor options. Ensure each distractor is plausible, homogenous in length and style, and covers common misconceptions. The correct answer should remain the same.\n\n{question}"), icon='psychology_alt').classes('text-xs')
                        ui.button("Predict Difficulty", on_click=lambda: handle_preset_click("Predict difficulty.", "Estimate the difficulty level (0.0-1.0) of this MCQ. Consider: clinical reasoning depth, distractor plausibility, vocabulary level, and whether it tests recall vs. application vs. analysis. Provide your rating with a brief justification.\n\n{question}"), icon='analytics').classes('text-xs')
                        ui.button("Convert to A2", on_click=lambda: handle_preset_click("Convert to A2.", "Convert this A1 (single-best-answer recall) question into an A2 (clinical case-based) question. Create a brief clinical vignette with patient demographics, presenting complaint, relevant history, examination findings, and/or investigation results. The answer options should test clinical reasoning rather than factual recall. Keep the same core content area.\n\n{question}"), icon='medication').classes('text-xs')
                        ui.button("Convert to A1", on_click=lambda: handle_preset_click("Convert to A1.", "Convert this A2 (case-based) question into an A1 (single-best-answer) question. Remove the clinical vignette and rephrase as a direct, concise question that tests the same core content through factual recall.\n\n{question}"), icon='functions').classes('text-xs')
                        ui.button("Typical Case", on_click=lambda: handle_preset_click("Make case more typical.", "If this is a case-based question, revise the clinical vignette to present a more classic, textbook-typical presentation of the condition. Ensure the case clearly points to the correct diagnosis through characteristic signs and symptoms. If it is already an A1 question, convert it to an A2 with a classic clinical presentation.\n\n{question}"), icon='biotech').classes('text-xs')

                    with ui.row().classes('w-full justify-start gap-2 p-2 flex-shrink-0') as img_presets:
                        async def handle_img_preset(face_text: str, prompt_tmpl: str):
                            actual = prompt_tmpl.format(prompt=question_data.get('image_prompt', 'None'))
                            await send_message(user_facing_text=face_text, actual_prompt=actual)
                        ui.button("Regenerate", on_click=lambda: handle_img_preset("Regenerate image.", "Generate a new image based on this prompt, with a different composition or style while maintaining the same medical content and educational purpose.\n\nOriginal prompt: {prompt}"), icon='restart_alt').classes('text-xs')
                        ui.button("Enhance", on_click=lambda: handle_img_preset("Enhance details.", "Enhance this image generation prompt. Add more specific details about: anatomical accuracy, pathological features, imaging modality characteristics, labeling requirements, and educational clarity. The image should be suitable for a medical examination question.\n\nOriginal prompt: {prompt}"), icon='enhance_photo_translate').classes('text-xs')

                    def update_btn_visibility(model_name: str):
                        if not model_name:
                            text_presets.set_visibility(True)
                            img_presets.set_visibility(False)
                            return
                        mc = None
                        for c in APP_STATE.get("saved_models_config", []):
                            if c.get('name') == model_name:
                                mc = c
                                break
                        is_img = mc and mc.get('model_type') in ['Image Generation', 'Multimodal'] if mc else False
                        text_presets.set_visibility(not is_img)
                        img_presets.set_visibility(is_img)
                    model_select.on_value_change(update_btn_visibility)
                    update_btn_visibility(model_select.value)

                with ui.card().classes('w-2/5 h-full'):
                    with ui.card_section().classes('w-full h-full overflow-y-auto'):
                        ui.markdown("**Current Question State**").classes('text-lg font-bold mb-2')
                        right_preview_panel = ui.column().classes('w-full')
            update_right_preview_panel(right_preview_panel)
            dialog.open()

    async def handle_generate():
        if not APP_STATE["saved_models_config"]:
            ui.notify("Error: No models configured.", color='negative')
            return
        model_names = {
            "writer": ui_refs["model_selects"]["Writer"].value,
            "reviewers": [
                ui_refs["model_selects"]["Reviewer 1 (Content)"].value,
                ui_refs["model_selects"]["Reviewer 2 (Structure)"].value,
                ui_refs["model_selects"]["Reviewer 3 (Format)"].value,
            ],
            "editor": ui_refs["model_selects"]["Editor/Final"].value,
            "image_generator": ui_refs["model_selects"]["Image Generation Model"].value,
        }
        if not all([model_names["writer"], model_names["editor"]] + model_names["reviewers"]):
            ui.notify("Error: Please assign models for all text roles.", color='negative')
            return
        is_img = bool(model_names["image_generator"])
        if is_img:
            ui.notify("Image generation enabled.", color='positive')
        else:
            ui.notify("No image model selected, text-only questions.", color='info')
        selected_kps = [
            {'keyword': ALL_NODES_BY_ID[kid]['label'].strip(), 'requirement': "None", 'count': int(data['number_input'].value)}
            for kid, data in APP_STATE["selected_kps_with_inputs"].items()
            if data['number_input'].value and int(data['number_input'].value) > 0
        ]
        if not selected_kps:
            ui.notify("Error: No syllabus items selected.", color='negative')
            return
        APP_STATE['is_generating'] = True
        APP_STATE['stop_requested'] = False
        status_text, question_preview_area, progress_row, step_progress, step_label, download_row = (ui_refs[k] for k in ["status_text", "question_preview_area", "progress_row", "step_progress", "step_label", "download_row"])
        status_text.value = "Preparing..."
        question_preview_area.clear()
        download_row.set_visibility(False)
        progress_row.set_visibility(True)
        with question_preview_area:
            ui.spinner(size='lg').classes('self-center mt-8')
        run_folder = BASE_OUTPUT_DIR / f"ExamRun_{datetime.datetime.now():%Y%m%d_%H%M%S}"
        reports_folder = run_folder / "individual_mcq_reports"
        run_folder.mkdir(parents=True, exist_ok=True)
        reports_folder.mkdir(exist_ok=True)
        ui_refs['run_folder'] = run_folder
        log_parts = []
        q_num, success_count = 0, 0
        total_q = sum(item['count'] for item in selected_kps)
        gen_q = []
        ui_refs['generated_questions'] = gen_q
        try:
            status_text.value = "Initializing AI model clients..."
            required = {model_names["writer"], model_names["editor"]} | set(model_names["reviewers"])
            if is_img:
                required.add(model_names["image_generator"])
            configs = [c for c in APP_STATE["saved_models_config"] if c['name'] in required]
            mcq_system = MCQDevelopmentSystem(configs)
            for item in selected_kps:
                for _ in range(item['count']):
                    if APP_STATE['stop_requested']:
                        raise StopRequestedError
                    q_num += 1
                    status_text.value = f"Generating: Question {q_num}/{total_q} ({item['keyword']})..."
                    history, q_img_path = [], None
                    prompt_ctx = generate_writer_prompt_context(item['keyword'], "None", APP_STATE["keyword_to_details_map"])
                    history, usage_stats, q_img_path, q_img_prompt = await run_full_mcq_pipeline(mcq_system, prompt_ctx, model_names, step_progress, step_label, q_num, is_img)
                    final_q = extract_final_question_from_history(history)
                    if final_q:
                        success_count += 1
                        gen_q.append({'number': success_count, 'text': final_q, 'image_path': q_img_path, 'image_prompt': q_img_prompt, 'chat_history': []})
                    report_html = generate_html_report_content(history, f"MCQ_{q_num}_Report", prompt_ctx, image_path=q_img_path)
                    save_content_to_file(report_html, reports_folder, f"mcq_{q_num}_report", "html")
                    log_parts.append(f"--- Question {q_num} ('{item['keyword']}') report generated ---")
            status_text.value = f"Done! Generated {success_count}/{total_q} questions."
        except StopRequestedError:
            status_text.value = f"Stopped by user. Completed {success_count} questions."
        except Exception as e:
            tb = traceback.format_exc()
            status_text.value = f"Error during generation: {e}"
            with question_preview_area:
                ui.html(f"<pre>{html.escape(tb)}</pre>")
        finally:
            progress_row.set_visibility(False)
            render_question_preview_area(gen_q)
            if gen_q:
                log_file = save_content_to_file("\n".join(log_parts), run_folder, "generation_summary_log", "txt")
                ui_refs['log_file_path'] = log_file
                update_questions()
                status_text.value += f" Files saved in {run_folder.resolve()}"
            APP_STATE['is_generating'] = False
            APP_STATE['stop_requested'] = False

    # --- Header ---
    with ui.row().classes('w-full items-center p-4 bg-blue-600 text-white'):
        ui.icon('biotech', size='2rem').classes('mr-2')
        ui.label('Multi-Agent Item Development (MAID)').classes('text-3xl font-bold')
        ui.element('div').classes('grow')

    ui.separator()

    # --- Main tabs ---
    with ui.tabs().classes('w-full') as main_tabs:
        mcq_tab = ui.tab('MCQ Development', icon='auto_stories')
        image_rag_tab = ui.tab('Image RAG', icon='photo_library')
        difficulty_tab = ui.tab('Difficulty Model Training', icon='model_training')

    with ui.tab_panels(main_tabs, value=mcq_tab).classes('w-full'):
        with ui.tab_panel(mcq_tab):
            with ui.row().classes('w-full no-wrap p-4 gap-4'):
                # Left: Item Generation Specification
                with ui.card().classes('w-1/3'):
                    ui.markdown("## Item Generation Specification")
                    with ui.expansion("Custom Syllabus Upload", icon='upload_file').classes('w-full').props('dense'):
                        ui.label("Upload a Markdown file with headings (# ## ### ...) to define your own syllabus tree. Bullet items (- or *) become selectable leaf nodes.").classes('text-xs text-gray-500')
                        syllabus_status = ui.label('').classes('text-xs')
                        def handle_syllabus_upload(e):
                            try:
                                md_bytes = e.content.read()
                                md_text = md_bytes.decode('utf-8')
                            except UnicodeDecodeError:
                                try:
                                    md_text = md_bytes.decode('gbk')
                                except Exception:
                                    ui.notify("Failed to decode file. Please use UTF-8 encoding.", type='negative')
                                    return
                            parsed = parse_markdown_syllabus(md_text)
                            normalized = normalize_syllabus_dict(parsed)
                            if not normalized:
                                ui.notify("No valid headings found in the Markdown file.", type='warning')
                                return
                            nonlocal tree_nodes, ALL_NODES_BY_ID
                            new_nodes = convert_to_tree_nodes(normalized)
                            tree_nodes = new_nodes
                            ALL_NODES_BY_ID = get_nodes_as_dict(new_nodes)
                            APP_STATE["keyword_to_details_map"] = flatten_syllabus_to_map(normalized)
                            APP_STATE["selected_kps_with_inputs"] = {}
                            tree.props({'nodes': new_nodes, 'ticked': []})
                            tree.update()
                            update_selection_area([])
                            leaf_count = sum(1 for n in ALL_NODES_BY_ID.values() if 'children' not in n)
                            syllabus_status.set_text(f"Loaded: {e.name} ({leaf_count} selectable items)")
                            syllabus_status.set_visibility(True)
                            ui.notify(f"Custom syllabus loaded: {e.name}", type='positive')
                        ui.upload(on_upload=handle_syllabus_upload, auto_upload=True, label="Drop .md file here or click to upload").props('accept=.md,.txt').classes('w-full mt-2')
                        with ui.row().classes('w-full mt-2'):
                            def reset_to_default():
                                nonlocal tree_nodes, ALL_NODES_BY_ID
                                tree_nodes = convert_to_tree_nodes(SYLLABUS_DATA_HIERARCHICAL)
                                ALL_NODES_BY_ID = get_nodes_as_dict(tree_nodes)
                                APP_STATE["keyword_to_details_map"] = flatten_syllabus_to_map(SYLLABUS_DATA_HIERARCHICAL)
                                APP_STATE["selected_kps_with_inputs"] = {}
                                tree.props({'nodes': tree_nodes, 'ticked': []})
                                tree.update()
                                update_selection_area([])
                                syllabus_status.set_text('')
                                ui.notify("Reset to default syllabus.", type='info')
                            ui.button("Reset to Default", on_click=reset_to_default, icon='restart_alt').props('flat dense size=sm color=grey')
                    search_input = ui.input(placeholder='Search syllabus...').props('clearable').classes('w-full')
                    tree = ui.tree(tree_nodes, label_key='label', tick_strategy='leaf', on_tick=lambda e: update_selection_area(e.value)).props('dense no-results-label="No matching items"')
                    search_input.bind_value(tree, 'filter')
                    ui.separator().classes('my-4')
                    ui.markdown("#### Selected Items & Question Count")
                    selection_area = ui.column().classes('w-full min-h-[100px] max-h-[250px] overflow-y-auto rounded border p-1 pb-6')
                    def remove_sel(item_id: str):
                        ticked = tree._props.get('ticked', [])
                        if item_id in ticked:
                            ticked.remove(item_id)
                            tree._props['ticked'] = ticked
                            tree.update()
                            update_selection_area(ticked)
                            ui.notify(f"Removed: {ALL_NODES_BY_ID[item_id]['label']}", color='info')
                    def update_selection_area(ticked_ids: List[str]):
                        selection_area.clear()
                        new_state = {}
                        with selection_area:
                            if not ticked_ids:
                                ui.label("Select items from the syllabus tree")
                            else:
                                for item_id in sorted(ticked_ids):
                                    if node := ALL_NODES_BY_ID.get(item_id):
                                        with ui.row().classes('w-full items-center no-wrap'):
                                            ui.button(icon='close', on_click=lambda _, i=item_id: remove_sel(i)).props('flat round dense size=sm').tooltip('Remove')
                                            ui.label(node['label']).classes('flex-grow')
                                            old = 1
                                            if prev := APP_STATE["selected_kps_with_inputs"].get(item_id):
                                                if pw := prev.get('number_input'):
                                                    if hasattr(pw, 'value') and pw.value is not None:
                                                        old = pw.value
                                            num_input = ui.number(label='Count', value=old, min=1, step=1).props('dense').classes('w-24')
                                            new_state[item_id] = {'number_input': num_input}
                        APP_STATE["selected_kps_with_inputs"] = new_state
                    def confirm_selection():
                        sels, total = [], 0
                        for kid, data in APP_STATE["selected_kps_with_inputs"].items():
                            ni = data.get('number_input')
                            if ni and hasattr(ni, 'value') and ni.value is not None and int(ni.value) > 0:
                                sels.append(f"Item: {kid} -> Count: {int(ni.value)}")
                                total += int(ni.value)
                        if not sels:
                            return ui.notify("No valid selections.", color='warning')
                        content = f"Selection confirmed - {datetime.datetime.now():%Y-%m-%d %H:%M:%S}\n\n" + "\n".join(sels) + f"\n\nTotal questions to generate: {total}"
                        saved = save_content_to_file(content, SELECTION_LOGS_DIR, "selection", "txt")
                        ui.notify(f"Selection saved. Will generate {total} questions.", color='positive')
                    def clear_all():
                        tree._props['ticked'] = []
                        tree.update()
                        update_selection_area([])
                        ui.notify("All selections cleared.", color='info')
                    with ui.row().classes('mt-4 w-full justify-start'):
                        ui.button("Confirm Selection", on_click=confirm_selection, color='primary')
                        ui.button("Clear All", on_click=clear_all, color='negative')

                # Middle: Generate
                with ui.card().classes('w-1/3'):
                    ui.markdown("## Generate Questions")
                    def request_stop():
                        APP_STATE['stop_requested'] = True
                        ui.notify("Stop requested! Will stop after current question.", color='warning', position='center', timeout=5000)
                    with ui.row().classes('w-full justify-center'):
                        ui_refs['generate_button'] = ui.button("Generate Item Sets", on_click=handle_generate, icon='auto_stories').classes('w-1/2').props('color=primary')
                        ui_refs['stop_button'] = ui.button("Stop Generation", on_click=request_stop, icon='stop_circle').classes('w-1/2').props('color=negative')
                    ui_refs['generate_button'].bind_enabled_from(APP_STATE, 'is_generating', backward=lambda x: not x)
                    ui_refs['stop_button'].bind_visibility_from(APP_STATE, 'is_generating')
                    ui_refs["progress_row"] = ui.row().classes('w-full items-center mt-2')
                    with ui_refs["progress_row"]:
                        ui_refs["step_progress"] = ui.linear_progress(value=0, show_value=False).classes('flex-grow')
                        ui_refs["step_label"] = ui.label('').classes('ml-2 text-xs')
                    ui_refs["progress_row"].set_visibility(False)
                    ui_refs["status_text"] = ui.textarea(label="Status").props('readonly outlined rows=2').classes('w-full mt-2')
                    ui_refs["question_preview_area"] = ui.column().classes('w-full border p-2 rounded').style('height: 400px; overflow-y: auto;')
                    ui_refs["download_row"] = ui.row().classes('w-full justify-around mt-4')
                    ui_refs["download_row"].set_visibility(False)

                # Right: Model Configuration
                with ui.card().classes('w-1/3'):
                    ui.markdown("## Model Configuration")
                    def update_model_ui():
                        configs = APP_STATE.get("saved_models_config", [])
                        text_choices = sorted([c['name'] for c in configs if c.get('model_type') == 'Text Output'])
                        img_choices = sorted([c['name'] for c in configs if c.get('model_type') in ['Image Generation', 'Multimodal']])
                        ui_refs["model_table"].rows = configs
                        ui_refs["model_table"].update()
                        for role, sel in ui_refs["model_selects"].items():
                            cur = sel.value
                            choices = img_choices if role == "Image Generation Model" else text_choices
                            sel.options = choices
                            sel.value = cur if cur in choices else (choices[0] if choices else None)
                            sel.update()
                    def save_model_config(name_in, key_in, url_in, type_sel, model_id_in, provider_sel):
                        name, key = name_in.value.strip(), key_in.value
                        mn = model_id_in.value.strip() or name
                        if not name or not key:
                            return ui.notify("Model name and API key required.", color='negative')
                        new_cfg = {"name": name, "model_name": mn, "api_key": key, "base_url": url_in.value.strip() or None, "model_type": type_sel.value, "provider": provider_sel.value}
                        APP_STATE["saved_models_config"] = [c for c in APP_STATE.get("saved_models_config", []) if c['name'] != name]
                        APP_STATE["saved_models_config"].append(new_cfg)
                        ui.notify(f"Config '{name}' saved.", color='positive')
                        name_in.value = key_in.value = url_in.value = model_id_in.value = ''
                        update_model_ui()
                    def delete_models():
                        if sel := ui_refs["model_table"].selected:
                            to_del = {row['name'] for row in sel}
                            APP_STATE["saved_models_config"] = [c for c in APP_STATE["saved_models_config"] if c['name'] not in to_del]
                            ui.notify(f"Deleted {len(to_del)} configs.", color='positive')
                            ui_refs["model_table"].selected.clear()
                            update_model_ui()
                    def load_presets(password: str):
                        if preset := PREDEFINED_KEYS.get(password):
                            loaded = []
                            for cfg in preset:
                                if all(k in cfg for k in ["name", "model_name", "api_key", "model_type"]):
                                    APP_STATE["saved_models_config"] = [c for c in APP_STATE["saved_models_config"] if c['name'] != cfg['name']]
                                    APP_STATE["saved_models_config"].append(cfg)
                                    loaded.append(cfg['name'])
                            if loaded:
                                update_model_ui()
                                ui.notify(f"Loaded {len(loaded)} presets.", color='positive', multi_line=True)
                        else:
                            ui.notify("Invalid password.", color='negative')
                    with ui.expansion("API & Model Settings", icon='settings', value=True).classes('w-full'):
                        with ui.card_section():
                            name_in = ui.input(label="Display Name")
                            provider_in = ui.select(['openai', 'nvidia'], label='Provider', value='openai')
                            model_id_in = ui.input(label="Model ID")
                            ak_in = ui.input(label="API Key", password=True)
                            bu_in = ui.input(label="Base URL")
                            mt_in = ui.select(['Text Output', 'Image Generation', 'Multimodal'], label='Model Type', value='Text Output')
                            ui.button("Save Config", on_click=lambda: save_model_config(name_in, ak_in, bu_in, mt_in, model_id_in, provider_in)).classes('w-full mt-4')
                        with ui.card_section():
                            pw_in = ui.input(label="Load Presets", placeholder="Enter preset password...")
                            ui.button("Load Presets", on_click=lambda: load_presets(pw_in.value)).classes('w-full mt-4')
                        ui.separator()
                        cols = [
                            {'name': 'name', 'label': 'Model Name', 'field': 'name', 'align': 'left'},
                            {'name': 'provider', 'label': 'Provider', 'field': 'provider', 'align': 'left'},
                            {'name': 'model_name', 'label': 'Model ID', 'field': 'model_name', 'align': 'left'},
                            {'name': 'model_type', 'label': 'Type', 'field': 'model_type', 'align': 'left'},
                            {'name': 'base_url', 'label': 'Base URL', 'field': 'base_url', 'align': 'left'}
                        ]
                        ui_refs["model_table"] = ui.table(columns=cols, rows=[], row_key='name', selection='multiple').classes('w-full h-40')
                        ui.button("Delete Selected", on_click=delete_models, color='negative', icon='delete')
                    with ui.card_section():
                        ui.markdown("#### Model Assignment")
                        text_roles = {
                            "Writer": "Initial draft and revisions",
                            "Reviewer 1 (Content)": "Medical content accuracy",
                            "Reviewer 2 (Structure)": "Question structure compliance",
                            "Reviewer 3 (Format)": "Language clarity and formatting",
                            "Editor/Final": "Synthesize feedback and final decision",
                        }
                        for role, info in text_roles.items():
                            with ui.row().classes('w-full items-center no-wrap'):
                                sel = ui.select([], label=role).classes('flex-grow').tooltip(info)
                                ui.button(icon='clear', on_click=lambda _, r=role, s=sel: sel.set_value(None)).props('flat dense size=sm')
                                ui_refs["model_selects"][role] = sel
                        ui.separator().classes('my-2')
                        with ui.row().classes('w-full items-center no-wrap'):
                            img_sel = ui.select([], label="Image Generation Model").classes('flex-grow').tooltip("Select an image generation model, leave empty for text-only")
                            ui.button(icon='clear', on_click=lambda: img_sel.set_value(None)).props('flat dense size=sm')
                            ui_refs["model_selects"]["Image Generation Model"] = img_sel
                        def random_assign():
                            t_models = [c['name'] for c in APP_STATE['saved_models_config'] if c['model_type'] == 'Text Output']
                            i_models = [c['name'] for c in APP_STATE['saved_models_config'] if c['model_type'] in ['Image Generation', 'Multimodal']]
                            if not t_models:
                                return ui.notify("No text models available.", color='warning')
                            for role, sel in ui_refs['model_selects'].items():
                                if role != "Image Generation Model":
                                    sel.set_value(random.choice(t_models))
                            if i_models and ui_refs['model_selects']["Image Generation Model"]:
                                ui_refs['model_selects']["Image Generation Model"].set_value(random.choice(i_models))
                        ui.button("Randomly Assign Models", on_click=random_assign, icon='shuffle').classes('w-full mt-4')

                        async def fill_demo_llms():
                            api_key = os.environ.get('NVIDIA_API_KEY', '').strip()
                            if not api_key:
                                ui.notify("Set NVIDIA_API_KEY before loading demo models.", color='negative')
                                return
                            ui.notify(f"Testing {len(NVIDIA_DEMO_MODELS)} NVIDIA models...", color='info', timeout=5000)
                            results = await check_nvidia_demo_models(api_key)
                            healthy = [result['config'] for result in results if result['healthy']]
                            failed = [result for result in results if not result['healthy']]
                            if not healthy:
                                details = '; '.join(f"{item['config']['model_name']}: {item['error']}" for item in failed)
                                ui.notify(f"No NVIDIA demo model is available. {details}", color='negative', multi_line=True, timeout=15000)
                                return
                            demo_names = {config['name'] for config in healthy}
                            APP_STATE['saved_models_config'] = [
                                config for config in APP_STATE.get('saved_models_config', [])
                                if config.get('name') not in demo_names
                            ] + healthy
                            update_model_ui()
                            text_roles = [role for role in ui_refs['model_selects'] if role != 'Image Generation Model']
                            assignments = assign_healthy_text_models([config['name'] for config in healthy], text_roles)
                            for role, model_name in assignments.items():
                                ui_refs['model_selects'][role].set_value(model_name)
                            ui.notify(
                                f"Loaded {len(healthy)} healthy NVIDIA models; {len(failed)} failed health checks.",
                                color='positive',
                                timeout=10000,
                            )

                        ui.button("Fill with Demo LLMs", on_click=fill_demo_llms, icon='bolt').classes('w-full mt-2')

        # Image RAG Tab
        with ui.tab_panel(image_rag_tab):
            with ui.row().classes('w-full no-wrap p-4 gap-4'):
                with ui.card().classes('w-1/2'):
                    ui.markdown("## Image Indexing").classes('text-xl font-bold mb-4')
                    def update_rag_stats():
                        stats = IMAGE_RAG_MANAGER.get_stats()
                        stats_label.text = f"Status: {'Enabled' if stats['enabled'] else 'Disabled'} | Images: {stats['total_images']} | Dim: {stats['dimension']}"
                    stats_label = ui.label("").classes('text-sm text-gray-600 mb-4')
                    update_rag_stats()
                    ui.separator().classes('my-4')
                    ui.label("Embedding Model:").classes('font-semibold')
                    rag_model_select = ui.select(['shibing624/text2vec-base-chinese', 'all-MiniLM-L6-v2', 'all-mpnet-base-v2'], value='shibing624/text2vec-base-chinese', label='Sentence Transformer Model').classes('w-full mb-4')
                    ui.label("Upload Images:").classes('font-semibold')
                    upload_container = ui.column().classes('w-full')
                    uploaded_files = []
                    with upload_container:
                        file_display = ui.label("No files selected").classes('text-gray-500')
                    def handle_image_upload(e):
                        uploaded_files.clear()
                        paths = []
                        for uf in e.files:
                            dest = IMAGE_RAG_DIR / uf.name
                            with open(dest, 'wb') as f:
                                f.write(uf.content.read())
                            paths.append(str(dest))
                            uploaded_files.append(str(dest))
                        if paths:
                            file_display.text = f"Selected: {', '.join([Path(p).name for p in paths])}"
                            ui.notify(f"Uploaded {len(paths)} image(s)", type='positive')
                    ui.upload(on_upload=handle_image_upload, multiple=True).props('accept=image/*').classes('w-full mb-2')
                    desc_input = ui.textarea(label="Image Descriptions (one per line, same order as uploads)", placeholder="Enter description for each image...").props('outlined autogrow').classes('w-full mb-4')
                    async def handle_index():
                        if not uploaded_files:
                            return ui.notify("Please upload images first.", type='warning')
                        descs_text = desc_input.value.strip()
                        if not descs_text:
                            return ui.notify("Please enter descriptions.", type='warning')
                        descs = [d.strip() for d in descs_text.split('\n') if d.strip()]
                        if len(descs) != len(uploaded_files):
                            return ui.notify(f"Mismatch: {len(uploaded_files)} images but {len(descs)} descriptions.", type='warning')
                        if rag_model_select.value != APP_STATE.get('rag_model_name'):
                            APP_STATE['rag_model_name'] = rag_model_select.value
                        ui.notify("Indexing...", type='info', timeout=10000)
                        result = await index_uploaded_images(uploaded_files, descs, rag_model_select.value)
                        if result['success']:
                            ui.notify(f"Indexed {result['success_count']}/{result['total_count']} images.", type='positive')
                            update_rag_stats()
                        else:
                            ui.notify(f"Error: {result.get('error')}", type='negative')
                        uploaded_files.clear()
                        file_display.text = "No files selected"
                    ui.button("Index Images", on_click=handle_index, icon='build').classes('w-full mb-4')
                    def clear_rag():
                        IMAGE_RAG_MANAGER.clear_all()
                        update_rag_stats()
                        ui.notify("Database cleared.", type='info')
                    ui.button("Clear Database", on_click=clear_rag, icon='delete', color='negative').classes('w-full')

                with ui.card().classes('w-1/2'):
                    ui.markdown("## RAG Search").classes('text-xl font-bold mb-4')
                    rag_switch = ui.switch("Enable RAG for Question Generation", value=APP_STATE.get('rag_search_enabled', False)).classes('w-full mb-4')
                    def update_rag():
                        APP_STATE['rag_search_enabled'] = rag_switch.value
                        ui.notify(f"RAG {'enabled' if rag_switch.value else 'disabled'}.", type='info')
                    rag_switch.on_value_change(update_rag)
                    ui.separator().classes('my-4')
                    ui.label("Search Query:").classes('font-semibold')
                    search_input = ui.input(label="Enter search query...", placeholder="e.g., ECG showing myocardial infarction").classes('w-full mb-2')
                    results_container = ui.column().classes('w-full border p-4 rounded min-h-[300px] max-h-[500px] overflow-y-auto')
                    async def handle_search():
                        query = search_input.value.strip()
                        if not query:
                            return ui.notify("Please enter a search query.", type='warning')
                        if IMAGE_RAG_MANAGER.get_stats().get('total_images', 0) == 0:
                            return ui.notify("No images in database.", type='warning')
                        results_container.clear()
                        with results_container:
                            ui.spinner(size='md')
                        results = search_rag_images(query, top_k=5)
                        results_container.clear()
                        with results_container:
                            if not results:
                                ui.label("No results found.").classes('text-gray-500')
                            else:
                                ui.markdown(f"**Found {len(results)} results:**").classes('mb-2')
                                for i, r in enumerate(results, 1):
                                    with ui.card().classes('mb-3'):
                                        with ui.row().classes('w-full items-center'):
                                            ui.badge(f"#{i}", color='primary').classes('mr-2')
                                            ui.markdown(f"**Distance:** {r['distance']:.2f}").classes('text-sm')
                                        img_p = Path(r['path'])
                                        if img_p.exists():
                                            try:
                                                rel = img_p.relative_to(BASE_OUTPUT_DIR)
                                                ui.image(f"/outputs/{rel}").classes('w-full max-w-md rounded-md mb-2')
                                            except ValueError:
                                                ui.image(str(img_p)).classes('w-full max-w-md rounded-md mb-2')
                                        ui.markdown(f"**Description:** {r['description']}").classes('text-sm')
                    ui.button("Search", on_click=handle_search, icon='search').classes('w-full')
                    ui.separator().classes('my-4')
                    ui.markdown("### RAG Integration Preview").classes('text-sm font-semibold mb-2')
                    ui.label("When enabled, relevant images will be automatically included in question generation prompts.").classes('text-xs text-gray-600')

        # Difficulty Model Training Tab
        with ui.tab_panel(difficulty_tab):
            with ui.row().classes('w-full no-wrap'):
                with ui.tabs().props('vertical').classes('w-32') as pred_tabs:
                    p_upload = ui.tab('Data Upload', icon='upload_file')
                    p_preprocess = ui.tab('Preprocessing', icon='cleaning_services')
                    p_features = ui.tab('Feature Engineering', icon='schema')
                    p_model = ui.tab('Model Training', icon='smart_toy')
                    p_predict = ui.tab('Difficulty Prediction', icon='online_prediction')
                    p_api = ui.tab('System Integration', icon='api')
                with ui.tab_panels(pred_tabs, value=p_upload).props('vertical').classes('flex-grow h-full') as panels_container:
                    create_difficulty_predictor_ui(
                        panels_container, p_upload, p_preprocess, p_features, p_model, p_predict, p_api, ui_refs
                    )


# ==============================================================================
# API Endpoints
# ==============================================================================
@app.post("/api/predict", response_model=PredictionResponse)
async def api_predict(request: PredictionRequest):
    STATE = APP_STATE["difficulty_predictor"]
    if not STATE.trained_model:
        raise HTTPException(status_code=503, detail="Model not trained or loaded.")
    try:
        texts = []
        for item in request.items:
            full = ""
            if isinstance(item, str):
                full = item
            elif isinstance(item, PredictionItem):
                full = item.stem
                if item.options:
                    for key in sorted(item.options.keys()):
                        full += f"\n{item.options[key]}"
                if item.answer:
                    full += f"\n{item.answer}"
            if STATE.preprocess_settings.get('lower'):
                full = full.lower()
            if STATE.preprocess_settings.get('punct'):
                full = re.sub(r'[^\w\s]', ' ', full)
            texts.append(full)
        if STATE.feature_method == 'tfidf':
            if not STATE.tfidf_vectorizer:
                raise HTTPException(status_code=500, detail="TF-IDF vectorizer not available.")
            fm = STATE.tfidf_vectorizer.transform(texts)
        elif STATE.feature_method == 'embedding':
            if not STATE.sentence_model_name:
                raise HTTPException(status_code=500, detail="Sentence model not available.")
            fm = await run.cpu_bound(_perform_encoding, texts, STATE.sentence_model_name)
        else:
            raise HTTPException(status_code=500, detail="Invalid feature method.")
        if STATE.stat_feature_names:
            sdf = pd.DataFrame({'text': texts})
            sv = []
            if 'char_count' in STATE.stat_feature_names:
                sv.append(sdf['text'].str.len())
            if 'word_count' in STATE.stat_feature_names:
                sv.append(sdf['text'].str.split().str.len())
            if 'sentence_count' in STATE.stat_feature_names:
                sv.append(sdf['text'].str.split(r'[.!?]').apply(lambda x: len([s for s in x if s.strip()])))
            if 'avg_word_length' in STATE.stat_feature_names:
                sv.append((sdf['text'].str.len() / sdf['text'].str.split().str.len().replace(0, 1)).fillna(0))
            sf = pd.concat(sv, axis=1).values
            fm = hstack([fm, sf]).tocsr() if isinstance(fm, csr_matrix) else np.hstack([fm, sf])
        preds = STATE.trained_model.predict(fm)
        return PredictionResponse(predictions=preds.tolist(), model_info=f"Model: {type(STATE.trained_model).__name__}, Features: {STATE.feature_method}")
    except Exception as e:
        log_difficulty_predictor(f"API error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/status")
def api_status():
    STATE = APP_STATE["difficulty_predictor"]
    return {"model_trained": STATE.trained_model is not None, "feature_method": STATE.feature_method, "sentence_model_name": STATE.sentence_model_name, "timestamp": time.time()}


# ==============================================================================
# Application Startup
# ==============================================================================
if __name__ in {"__main__", "__mp_main__"}:
    BASE_OUTPUT_DIR.mkdir(exist_ok=True)
    SELECTION_LOGS_DIR.mkdir(exist_ok=True)
    IMAGES_DIR.mkdir(exist_ok=True)
    IMAGE_RAG_DIR.mkdir(parents=True, exist_ok=True)
    IMAGE_RAG_MANAGER = ImageRAGManager(IMAGE_RAG_DIR)
    APP_STATE["image_rag"] = IMAGE_RAG_MANAGER
    log_difficulty_predictor('MAID Platform started successfully')
    print(f"Output directory: '{BASE_OUTPUT_DIR.resolve()}'")
    if FAISS_AVAILABLE:
        print("Image RAG module enabled with FAISS support")
    else:
        print("Image RAG module disabled (FAISS not installed)")
    ui.run(title="Multi-Agent Item Development (MAID)", dark=False, host="0.0.0.0", port=8080, reload=False, uvicorn_logging_level='warning')
